import pyvisa as visa
import time
from customtkinter import *
import pandas as pd
from CTkMessagebox import *
from datetime import datetime
from threading import Thread, Event
from docx import Document
import os
import docx2pdf

class APTracker:
    def __init__(self, name):
        self.name = name
        self.ip_ps = 'TCPIP0::192.168.10.11::5025::SOCKET'
        self.ip_zna = 'TCPIP0::192.168.10.10::hislip0::INSTR'
        self.ip_lo = 'TCPIP0::192.168.10.3::inst0::INSTR'

        self.stop_event = Event()

        self.sfb1_sl = "123"
        self.sfb2_sl = "123"
        self.sfb3_sl = "123"
        self.hmrx_sl = "123"
        self.qsrx_sl = "123"

    def folder_checker(self):
        rm = visa.ResourceManager()
        device = rm.open_resource(self.ip_zna)

        def folder_maker(parent,child):
            device.write("MMEMory:CDIRectory '%s'" % (parent))
            time.sleep(0.1)
            response = device.query('MMEM:CAT? "%s"' % (parent))
            if child.lower()  not in response:
                # Create the folder on the ZNA
                device.write('MMEM:MDIR "%s"' % (child))

        folder_maker("C:","Synergy")
        level_1 = ["QSRx","QSRx_int","SFB1","SFB2","SFB3"]
        level_2 = ["RF","BITE"]
        for check1 in level_1:
            folder_maker("C:\\Synergy",check1)
            if check1 != "QSRx":
                for check2 in level_2:
                    folder_maker("C:\\Synergy\\%s" % (check1),check2)
        
        device.close()
        rm.close()

    def turn_off_ps1(self):
        ip_ps = self.ip_ps

        rm = visa.ResourceManager()
        device = rm.open_resource(ip_ps)

        device.write('OUTP:GEN 0')
        time.sleep(0.1)
        device.write("INST:NSEL 4")
        time.sleep(0.1)
        device.write("OUTP:SEL 0")
        time.sleep(0.1)
        device.write("INST:NSEL 3")
        time.sleep(0.1)
        device.write("OUTP:SEL 0")
        time.sleep(0.1)
        device.write("INST:NSEL 2")
        time.sleep(0.1)
        device.write("OUTP:SEL 0")
        time.sleep(0.1)
        device.write("INST:NSEL 1")
        time.sleep(0.1)
        device.write("OUTP:SEL 0")
        time.sleep(0.1)

        device.close()
        rm.close()

    def ps_output(self,state=0):
            ip_ps = self.ip_ps

            rm = visa.ResourceManager()
            device = rm.open_resource(ip_ps)

            device.write('OUTP:GEN %s' % (state))
            
            device.close()
            rm.close()

    def mbox(self,message):
        box = CTkMessagebox(title='Instrument Control', message=message, option_1='OK',option_2='Abort',icon_size=(40,40),width=500,justify='centre',wraplength=500,font=("Arial",15))        
        response = box.get()
        return response

    def stop_task(self):
        self.stop_event.set()
        self.ps_output(0)
        self.turn_off_ps1()
        self.rf_off()
        _ = self.mbox("Process has been aborted!")

    def resource_path(self,relative_path):
        try:
            base_path = sys._MEIPASS
        except Exception:
            base_path = os.path.abspath(".")

        return os.path.join(base_path, relative_path)

    def report_qsr(self,folder,template,output,pdf_output,new_value):
        """
        Edit a specific cell in the table of a Word document.

        Args:
            word_file (str): Path to the input Word document.
            output_file (str): Path to save the updated Word document.
            new_value (str): New value to put in the cell.
        """
        os.mkdir("test_data/%s" % folder)
        # Load the Word document
        doc1 = Document(self.resource_path(template))
        
        table1 = doc1.tables[1]  # Access the first table

        for i in range(0,27):
        # Edit the cell value
            p2 = table1.cell(i+3, 5)
            p4 = table1.cell(i+3, 6)
            pha = table1.cell(i+3, 7)
            amp = table1.cell(i+3, 8)
            
            p2.text = str(round(new_value.iloc[i,5],2))
            p4.text = str(round(new_value.iloc[i,6],2))
            pha.text = str(round(new_value.iloc[i,7],2))
            amp.text = str(round(new_value.iloc[i,8],2))

        
        table2 = doc1.tables[3]

        for j in range(0,11):
        # Edit the cell value
            p2 = table2.cell(j, 5)
            p4 = table2.cell(j, 6)
            pha = table2.cell(j, 7)
            amp = table2.cell(j, 8)
            
            p2.text = str(round(new_value.iloc[j+27,5],2))
            p4.text = str(round(new_value.iloc[j+27,6],2))
            pha.text = str(round(new_value.iloc[j+27,7],2))
            amp.text = str(round(new_value.iloc[j+27,8],2))

        # Save the updated document
        doc1.save("test_data/%s/%s" % (folder,output))
        
        docx2pdf.convert("test_data/%s/%s" % (folder,output),"test_data/%s/%s" % (folder,pdf_output))

    def report_sfb(self,folder,template,output,pdf_output,new_value):
        
        # Load the Word document
        doc1 = Document(self.resource_path(template))
        
        table1 = doc1.tables[1]  # Access the first table

        for i in range(4):
        # Edit the cell value
            amp = table1.cell(i+1, 2)
            pha = table1.cell(i+1, 3)
            
            pha.text = str(round(new_value.iloc[i,2],2))
            amp.text = str(round(new_value.iloc[i,1],2))

        # Save the updated document
        doc1.save("test_data/%s/%s" % (folder,output))
        
        docx2pdf.convert("test_data/%s/%s" % (folder,output),"test_data/%s/%s" % (folder,pdf_output))

    def rf_off(self):
        rm = visa.ResourceManager()
        device_vna = rm.open_resource(self.ip_zna)
        device_vna.write('OUTP OFF')
        device_vna.close()
        rm.close()

        ip_lo = self.ip_lo

        rm = visa.ResourceManager()
        device_lo = rm.open_resource(ip_lo)
        device_lo.write(':OUTPut:STATe %d' % (0))
        device_lo.close()
        rm.close()
    
    ####SFB Module####

    def submit_ps1_sfb(self):
        ip_ps = self.ip_ps

        rm = visa.ResourceManager()
        device = rm.open_resource(ip_ps)

        device.write('*RST')

        device.write("INST OUT1")
        time.sleep(0.1)
        device.write('OUTP:SEL 1')
        device.write('VOLT:PROT 15')
        device.write('VOLT 9')
        device.write('CURR 1')

        time.sleep(0.1)

        device.write("INST OUT2")
        time.sleep(0.1)
        device.write('OUTP:SEL 1')
        device.write('VOLT:PROT 15')
        device.write('VOLT 9')
        device.write('CURR 6')

        time.sleep(0.1)

        device.write("INST OUT4")
        time.sleep(0.1)
        device.write('OUTP:SEL 1')
        device.write('VOLT:PROT 7.5')
        device.write('VOLT 5')
        device.write('CURR 0.3')
        time.sleep(0.1)
        device.write('OUTP 1')

        CTkMessagebox(title='Warning!', message='Please wait for 40 seconds before continuing.', option_1='OK')

    def sfb_settings(self,bands,window):
        ip_zna = self.ip_zna
        start_f  = bands[0]
        stop_f = bands[1]
        power=-20
        sweep_pts=101
        if_bw=100

        rm = visa.ResourceManager()
        device_vna = rm.open_resource(ip_zna)

        device_vna.write(':DISPlay:WINDow%s:STATe %d' % (window,1))
        device_vna.write(':SENSe%s:FREQuency:STARt %G' % (window,start_f*1e9))
        device_vna.write(':SENSe%s:FREQuency:STOP %G' % (window,stop_f*1e9))
        device_vna.write(':SOURce:POWer:LEVel:IMMediate:AMPLitude %G' % (power))
        device_vna.write(':SENSe%s:SWEep:POINts %d' % (window,sweep_pts))
        device_vna.write(':SENSe%s:BANDwidth:RESolution %G' % (window,if_bw*1e3))
        device_vna.write(':OUTPut:STATe %d' % (1))
        time.sleep(1)
        device_vna.close()
        rm.close()
    
    def phase_ref_sfb(self,band_num,sub_band,window=2):
        ip_zna = self.ip_zna
        bands = [[[0.5,0.8],[0.7,1.2],[1.1,1.6],[1.5,2.2]],
                [[2.0,2.75],[2.5,3.75],[3.5,5.25],[5,6.25]],
                [[6.0,9.5],[9.0,13.0],[12.5,14.5],[14.0,18.0]]]
        
        current = bands[band_num-1][sub_band]
        
        self.sfb_settings(current,window=window)

        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)
        device.write(":CALC%s:PAR:SDEF 'Phase%s', 'S21'" % (window,window))
        time.sleep(1)
        device.write("DISPlay:WINDow%s:STATe ON" % (window))
        device.write(":DISP:WIND%s:TITL:DATA 'Phase: Band%s'" % (window,band_num))
        device.write(":DISP:WIND%s:TRAC:EFEED 'Phase%s'" % (window,window))
        time.sleep(1)
        device.write(":DISPlay:WINDow%s:TRACe:Y:PDIVision 5,'Phase%s'" % (window,window))
        device.write("DISP:WIND%s:TRAC:Y:RLEV 0, 'Phase%s'" % (window,window))
        device.write("DISP:WIND%s:TRAC:Y:RPOS 50, 'Phase%s'" % (window,window))
        device.write(':CALCulate%s:FORMat %s' % (window,'PHAse'))
        time.sleep(0.5)
        device.write(":TRAC:COPY:MATH 'Phase%s_mem','Phase%s'" % (window,window))
        device.write(":DISP:WIND%s:TRAC:EFEED 'Phase%s_mem'" % (window,window))
        time.sleep(1)
        device.write(":CALC%s:MATH:SDEF 'Phase%s / Phase%s_mem'" % (window,window,window))
        device.write(':CALC%s:MATH:STAT ON' % (window))
        time.sleep(1)
        device.write("TRAC:COPY:MATH 'Phase%s_Ch1','Phase%s'" % (window,window))
        time.sleep(1)
        #device.write(":CALC:MATH:STAT OFF")
        device.write(":DISP:WIND%s:TRAC:EFEED 'Phase%s_Ch1'" % (window,window))
        device.write(":DISPlay:WINDow%s:TRACe:Y:PDIVision 5,'Phase%s_Ch1'" % (window,window))
        device.write("DISP:WIND%s:TRAC:Y:RLEV 0, 'Phase%s_Ch1'" % (window,window))
        device.write("DISP:WIND%s:TRAC:Y:RPOS 50, 'Phase%s_Ch1'" % (window,window))

        device.write("DISP:TRAC:SHOW 'Phase%s_mem', OFF" % (window))

        device.close()
        rm.close()
    
    def gain_ref(self,band_num,sub_band,type='sfb',window=2):

        ip_zna = self.ip_zna

        bands = [[[0.5,0.8],[0.7,1.2],[1.1,1.6],[1.5,2.2]],
                [[2.0,2.75],[2.5,3.75],[3.5,5.25],[5,6.25]],
                [[6.0,9.5],[9.0,13.0],[12.5,14.5],[14.0,18.0]]]
        
        current = bands[band_num-1][sub_band]
        if type =='sfb':
            self.sfb_settings(current,window=window)

        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)

        if window==1:
            device.write(":CONFigure:TRACe1:REName 'Trc1', 'Gain%s'" % (window))
        else:
            device.write(":CALC%s:PAR:SDEF 'Gain%s', 'S21'" % (window,window))
            device.write(":DISP:WIND%s:TRAC:EFEED 'Gain%s'" % (window,window))
            device.write(":DISPlay:WINDow%s:TRACe:Y:PDIVision 1,'Gain%s'" % (window,window))
        
        #device.write("CONFigure:CHANnel2:TRACe:REName 'Gain'")
        device.write("DISP:WIND%s:TRAC:Y:RLEV 0, 'Gain%s'" % (window,window))
        device.write("DISP:WIND%s:TRAC:Y:RPOS 50, 'Gain%s'" % (window,window))
        device.write("DISPlay:WINDow%s:STATe ON" % (window))
        device.write(":DISP:WIND%s:TITL:DATA 'Gain: Band%s'" % (window,band_num))
        device.write(':CALCulate%s:FORMat %s' % (window,'MLOGarithmic'))
        time.sleep(0.5)
        device.write(":TRAC:COPY:MATH 'Gain%s_mem','Gain%s'" % (window,window))
        device.write(":DISP:WIND%s:TRAC:EFEED 'Gain%s_mem'" % (window,window))
        time.sleep(0.5)
        device.write(":CALC%s:MATH:SDEF 'Gain%s / Gain%s_mem'" % (window,window,window))
        device.write(':CALC%s:MATH:STAT ON' % (window))
        time.sleep(0.5)
        device.write("TRAC:COPY:MATH 'Gain%s_Ch1','Gain%s'" % (window,window))
        time.sleep(0.5)
        #device.write(":CALC:MATH:STAT OFF")
        device.write(":DISP:WIND%s:TRAC:EFEED 'Gain%s_Ch1'" % (window,window))
        device.write(":DISPlay:WINDow%s:TRACe:Y:PDIVision 1,'Gain%s_Ch1'" % (window,window))
        device.write("DISP:WIND%s:TRAC:Y:RLEV 0, 'Gain%s_Ch1'" % (window,window))
        device.write("DISP:WIND%s:TRAC:Y:RPOS 50, 'Gain%s_Ch1'" % (window,window))
        device.write("DISP:TRAC:SHOW 'Gain%s_mem', OFF" % (window))
        time.sleep(0.5)

        device.close()
        rm.close()
    
    def new_phase(self,window,p_num=2):
        ip_zna = self.ip_zna

        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)

        device.write(":CALCulate%s:PARameter:SELect 'Phase%s'" % (window,window))
        device.write(":CALC%s:MATH:SDEF 'Phase%s / Phase%s_mem'" % (window,window,window))
        device.write(':CALC%s:MATH:STAT ON' % (window))
        device.write("TRAC:COPY:MATH 'Phase%s_Ch%s','Phase%s'" % (window,p_num,window))
        time.sleep(0.5)
        #device.write(":CALC:MATH:STAT OFF")
        #device.write("DISPlay:WINDow2:STATe ON")
        device.write(":DISP:WIND%s:TRAC:EFEED 'Phase%s_Ch%s'" % (window,window,p_num))
        time.sleep(0.5)
        device.write(":DISPlay:WINDow%s:TRACe:Y:PDIVision 5,'Phase%s_Ch%s'" % (window,window,p_num))
        device.write("DISP:WIND%s:TRAC:Y:RLEV 0, 'Phase%s_Ch%s'" % (window,window,p_num))
        device.write("DISP:WIND%s:TRAC:Y:RPOS 50, 'Phase%s_Ch%s'" % (window,window,p_num))

        time.sleep(0.5)

        device.close()
        rm.close()
    
    def new_gain(self,window,g_num):
        ip_zna = self.ip_zna
        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)

        
        device.write(":CALCulate%s:PARameter:SELect 'Gain%s'" % (window,window))
        device.write(":CALC%s:MATH:SDEF 'Gain%s / Gain%s_mem'" % (window,window,window))
        device.write(':CALC%s:MATH:STAT ON' % (window))
        device.write("TRAC:COPY:MATH 'Gain%s_Ch%s','Gain%s'" % (window,g_num,window))
        time.sleep(0.5)
        #device.write(":CALC:MATH:STAT OFF")
        #device.write("DISPlay:WINDow2:STATe ON")
        device.write(":DISP:WIND%s:TRAC:EFEED 'Gain%s_Ch%s'" % (window,window,g_num))
        time.sleep(0.5)
        device.write(":DISPlay:WINDow%s:TRACe:Y:PDIVision 1,'Gain%s_Ch%s'" % (window,window,g_num))
        device.write("DISP:WIND%s:TRAC:Y:RLEV 0, 'Gain%s_Ch%s'" % (window,window,g_num))
        device.write("DISP:WIND%s:TRAC:Y:RPOS 50, 'Gain%s_Ch%s'" % (window,window,g_num))

        device.close()
        rm.close()
    
    def track_sfb1(self):

        ip_zna = self.ip_zna

        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()

        band_ranges = [" (0.5 GHz to 0.8 GHz)",
                    " (0.7 GHz to 1.2 GHz)",
                    " (1.1 GHz to 1.6 GHz)",
                    " (1.5 GHz to 2.2 GHz)"]
        
        for i in range(4):
            if self.stop_event.is_set():
                return
            band = 'Change to Band' + str(i+1) + band_ranges[i]
            f_band = self.mbox(band)
            if f_band == 'OK':
                self.gain_ref(band_num=1,sub_band=i,type='sfb',window=2*i+1)
                time.sleep(1)
                self.phase_ref_sfb(band_num=1,sub_band=i,window=2*i+2)
            else:
                self.stop_task()
            
        for j in range(3):
            if self.stop_event.is_set():
                return
            mess = 'Change to Channel ' + str(j+2) + '\nChange to Band 1' + band_ranges[0]
            ref_port = self.mbox(mess)
            if ref_port == 'OK':
                self.new_gain(1,j+2)
                self.new_phase(2,j+2)

                f_band = self.mbox('Change to Band 2 %s' % (band_ranges[1]))
                if f_band == 'OK':
                    self.new_gain(3,j+2)
                    self.new_phase(4,j+2)
                
                    f_band = self.mbox('Change to Band 3 %s' % (band_ranges[2]))
                    if f_band == 'OK':
                        self.new_gain(5,j+2)
                        self.new_phase(6,j+2)
                    
                        f_band = self.mbox('Change to Band 4 %s' % (band_ranges[3]))
                        if f_band == 'OK':
                            self.new_gain(7,j+2)
                            self.new_phase(8,j+2)
                        else:
                            self.stop_task()
                    else:
                        self.stop_task()
                else:
                    self.stop_task()
            else:
                self.stop_task()

        time.sleep(0.5)
        self.marker_sfb(1,False)


        self.ps_output(0)
        dut = self.sfb1_sl
        self.save_diagram("SFB1","RF",dut)
    
    def track_sfb2(self):
        ip_zna = self.ip_zna
        
        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()

        band_ranges = [" (2.0 GHz to 2.75 GHz)",
                    " (2.5 GHz to 3.75 GHz)",
                    " (3.5 GHz to 5.25 GHz)",
                    " (5.0 GHz to 6.25 GHz)"]
        
        for i in range(4):
            if self.stop_event.is_set():
                return
            band = 'Change to Band' + str(i+1) + band_ranges[i]
            f_band = self.mbox(band)
            if f_band == 'OK':
                self.gain_ref(band_num=2,sub_band=i,type='sfb',window=2*i+1)
                self.phase_ref_sfb(band_num=2,sub_band=i,window=2*i+2)
            
        for j in range(3):
            if self.stop_event.is_set():
                return
            mess = 'Change to Channel ' + str(j+2) + '\nChange to Band 1' + band_ranges[0]
            ref_port = self.mbox(mess)
            if ref_port == 'OK':
                self.new_gain(1,j+2)
                self.new_phase(2,j+2)

                f_band = self.mbox('Change to Band 2 %s' % (band_ranges[1]))
                if f_band == 'OK':
                    self.new_gain(3,j+2)
                    self.new_phase(4,j+2)
                
                    f_band = self.mbox('Change to Band 3 %s' % (band_ranges[2]))
                    if f_band == 'OK':
                        self.new_gain(5,j+2)
                        self.new_phase(6,j+2)
                    
                        f_band = self.mbox('Change to Band 4 %s' % (band_ranges[3]))
                        if f_band == 'OK':
                            self.new_gain(7,j+2)
                            self.new_phase(8,j+2)
                        else:
                            self.stop_task()
                    else:
                        self.stop_task()
                else:
                    self.stop_task()
            else:
                self.stop_task()
        
        time.sleep(0.5)
        self.marker_sfb(2,False)
        self.ps_output(0)
        dut = self.sfb2_sl
        self.save_diagram("SFB2","RF",dut)
    
    def track_sfb3(self):
        ip_zna = self.ip_zna
        
        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()
        
        
        band_ranges = [" (6.0 GHz to 9.5 GHz)",
                    " (9.0 GHz to 13.0 GHz)",
                    " (12.5 GHz to 14.5 GHz)",
                    " (14.0 GHz to 18.0 GHz)"]

        for i in range(4):
            if self.stop_event.is_set():
                return
            band = 'Change to Band' + str(i+1) + band_ranges[i]
            f_band = self.mbox(band)
            if f_band == 'OK':
                self.gain_ref(band_num=3,sub_band=i,type='sfb',window=2*i+1)
                self.phase_ref_sfb(band_num=3,sub_band=i,window=2*i+2)
            
        for j in range(3):
            if self.stop_event.is_set():
                return
            mess = 'Change to Channel ' + str(j+2) + '\nChange to Band 1' + band_ranges[0]
            ref_port = self.mbox(mess)
            if ref_port == 'OK':
                self.new_gain(1,j+2)
                self.new_phase(2,j+2)

                f_band = self.mbox('Change to Band 2 %s' % (band_ranges[1]))
                if f_band == 'OK':
                    self.new_gain(3,j+2)
                    self.new_phase(4,j+2)
                
                    f_band = self.mbox('Change to Band 3 %s' % (band_ranges[2]))
                    if f_band == 'OK':
                        self.new_gain(5,j+2)
                        self.new_phase(6,j+2)
                    
                        f_band = self.mbox('Change to Band 4 %s' % (band_ranges[3]))
                        if f_band == 'OK':
                            self.new_gain(7,j+2)
                            self.new_phase(8,j+2)
                        else:
                            self.stop_task()
                    else:
                        self.stop_task()
                else:
                    self.stop_task()
            else:
                self.stop_task()
        
        time.sleep(0.5)
        self.marker_sfb(3,False)
        self.ps_output(0)
        dut = self.sfb3_sl
        self.save_diagram("SFB3","RF",dut)
    
    def track_sfb1_bite(self):
        ip_zna = self.ip_zna
        
        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()
        
        
        band_ranges = [" (0.5 GHz to 0.8 GHz)",
                    " (0.7 GHz to 1.2 GHz)",
                    " (1.1 GHz to 1.6 GHz)",
                    " (1.5 GHz to 2.2 GHz)"]
        
        bite_set = self.mbox('Set to BITE mode')
        
        if bite_set == 'OK':
            for i in range(4):
                if self.stop_event.is_set():
                    return
                band = 'Change to Band' + str(i+1) + band_ranges[i]
                f_band = self.mbox(band)
                if f_band == 'OK':
                    self.gain_ref(band_num=1,sub_band=i,type='sfb',window=2*i+1)
                    self.phase_ref_sfb(band_num=1,sub_band=i,window=2*i+2)
                
            for j in range(3):
                if self.stop_event.is_set():
                    return
                mess = 'Change to Channel ' + str(j+2) + '\nChange to Band 1' + band_ranges[0]
                ref_port = self.mbox(mess)
                if ref_port == 'OK':
                    self.new_gain(1,j+2)
                    self.new_phase(2,j+2)

                    f_band = self.mbox('Change to Band 2 %s' % (band_ranges[1]))
                    if f_band == 'OK':
                        self.new_gain(3,j+2)
                        self.new_phase(4,j+2)
                    
                        f_band = self.mbox('Change to Band 3 %s' % (band_ranges[2]))
                        if f_band == 'OK':
                            self.new_gain(5,j+2)
                            self.new_phase(6,j+2)
                        
                            f_band = self.mbox('Change to Band 4 %s' % (band_ranges[3]))
                            if f_band == 'OK':
                                self.new_gain(7,j+2)
                                self.new_phase(8,j+2)
                            else:
                                self.stop_task()
                        else:
                            self.stop_task()
                    else:
                        self.stop_task()
                else:
                    self.stop_task()

            time.sleep(0.5)
            self.marker_sfb(1,True)
            self.ps_output(0)
            dut = self.sfb1_sl
            self.save_diagram("SFB1","BITE",dut)       
    
    def track_sfb2_bite(self):
        ip_zna = self.ip_zna
        
        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()
        
        
        band_ranges = [" (2.0 GHz to 2.75 GHz)",
                    " (2.5 GHz to 3.75 GHz)",
                    " (3.5 GHz to 5.25 GHz)",
                    " (5.0 GHz to 6.25 GHz)"]
        
        bite_set = self.mbox('Set to BITE mode')
        
        if bite_set == 'OK':
            for i in range(4):
                if self.stop_event.is_set():
                    return
                band = 'Change to Band' + str(i+1) + band_ranges[i]
                f_band = self.mbox(band)
                if f_band == 'OK':
                    self.gain_ref(band_num=2,sub_band=i,type='sfb',window=2*i+1)
                    self.phase_ref_sfb(band_num=2,sub_band=i,window=2*i+2)
                
            for j in range(3):
                if self.stop_event.is_set():
                    return
                mess = 'Change to Channel ' + str(j+2) + '\nChange to Band 1' + band_ranges[0]
                ref_port = self.mbox(mess)
                if ref_port == 'OK':
                    self.new_gain(1,j+2)
                    self.new_phase(2,j+2)

                    f_band = self.mbox('Change to Band 2 %s' % (band_ranges[1]))
                    if f_band == 'OK':
                        self.new_gain(3,j+2)
                        self.new_phase(4,j+2)
                    
                        f_band = self.mbox('Change to Band 3 %s' % (band_ranges[2]))
                        if f_band == 'OK':
                            self.new_gain(5,j+2)
                            self.new_phase(6,j+2)
                        
                            f_band = self.mbox('Change to Band 4 %s' % (band_ranges[3]))
                            if f_band == 'OK':
                                self.new_gain(7,j+2)
                                self.new_phase(8,j+2)
                            else:
                                self.stop_task()
                        else:
                            self.stop_task()
                    else:
                        self.stop_task()
                else:
                    self.stop_task()
            
            time.sleep(0.5)
            self.marker_sfb(2,True)
            self.ps_output(0)
            dut = self.sfb2_sl
            self.save_diagram("SFB2","BITE",dut)
    
    def track_sfb3_bite(self):
        ip_zna = self.ip_zna
        
        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()
        
        
        band_ranges = [" (6.0 GHz to 9.5 GHz)",
                    " (9.0 GHz to 13.0 GHz)",
                    " (12.5 GHz to 14.5 GHz)",
                    " (14.0 GHz to 18.0 GHz)"]

        bite_set = self.mbox('Set to BITE mode')
        
        if bite_set == 'OK':
            for i in range(4):
                if self.stop_event.is_set():
                    return
                band = 'Change to Band' + str(i+1) + band_ranges[i]
                f_band = self.mbox(band)
                if f_band == 'OK':
                    self.gain_ref(band_num=3,sub_band=i,type='sfb',window=2*i+1)
                    self.phase_ref_sfb(band_num=3,sub_band=i,window=2*i+2)
                
            for j in range(3):
                if self.stop_event.is_set():
                    return
                mess = 'Change to Channel ' + str(j+2) + '\nChange to Band 1' + band_ranges[0]
                ref_port = self.mbox(mess)
                if ref_port == 'OK':
                    self.new_gain(1,j+2)
                    self.new_phase(2,j+2)

                    f_band = self.mbox('Change to Band 2 %s' % (band_ranges[1]))
                    if f_band == 'OK':
                        self.new_gain(3,j+2)
                        self.new_phase(4,j+2)
                    
                        f_band = self.mbox('Change to Band 3 %s' % (band_ranges[2]))
                        if f_band == 'OK':
                            self.new_gain(5,j+2)
                            self.new_phase(6,j+2)
                        
                            f_band = self.mbox('Change to Band 4 %s' % (band_ranges[3]))
                            if f_band == 'OK':
                                self.new_gain(7,j+2)
                                self.new_phase(8,j+2)
                            else:
                                self.stop_task()
                        else:
                            self.stop_task()
                    else:
                        self.stop_task()
                else:
                    self.stop_task()
            
            time.sleep(0.5)
            self.marker_sfb(3,True)
            self.ps_output(0)
            dut = self.sfb1_sl
            self.save_diagram("SFB3","BITE",dut)
    
    def marker_sfb(self,band_num,bite=False):
        rm = visa.ResourceManager()
        device = rm.open_resource(self.ip_zna)

        #device.write('*RST')
        amp_track = []
        phase_track = []

        for i in [1,3,5,7]:
            for j in range(4):
                device.write(":CALCulate%s:PARameter:SELect 'Gain%s_Ch%s'" % (i,i,j+1))
                device.write(":CALCulate%s:MARKer%s ON" % (i,2*j+1))
                device.write(":CALCulate%s:MARKer%s:MAX" % (i,2*j+1))

                device.write(":CALCulate%s:MARKer%s ON" % (i,2*j+2))
                device.write(":CALCulate%s:MARKer%s:MIN" % (i,2*j+2))

        for m in [2,4,6,8]:
            for n in range(4):
                device.write(":CALCulate%s:PARameter:SELect 'Phase%s_Ch%s'" % (m,m,n+1))
                device.write(":CALCulate%s:MARKer%s ON" % (m,2*n+1))
                device.write(":CALCulate%s:MARKer%s:MAX" % (m,2*n+1))

                device.write(":CALCulate%s:MARKer%s ON" % (m,2*n+2))
                device.write(":CALCulate%s:MARKer%s:MIN" % (m,2*n+2))


        for i in [1,3,5,7]:
            device.write(":CALCulate%s:PARameter:SELect 'Gain%s_Ch1'" % (i,i))
            max_val1 = float(device.query("CALC%s:MARK1:Y?" % (i)))
            min_val1 = float(device.query("CALC%s:MARK2:Y?" % (i)))

            device.write(":CALCulate%s:PARameter:SELect 'Gain%s_Ch2'" % (i,i))
            max_val2 = float(device.query("CALC%s:MARK3:Y?" % (i)))
            min_val2 = float(device.query("CALC%s:MARK4:Y?" % (i)))

            device.write(":CALCulate%s:PARameter:SELect 'Gain%s_Ch3'" % (i,i))
            max_val3 = float(device.query("CALC%s:MARK5:Y?" % (i)))
            min_val3 = float(device.query("CALC%s:MARK6:Y?" % (i)))

            device.write(":CALCulate%s:PARameter:SELect 'Gain%s_Ch4'" % (i,i))
            max_val4 = float(device.query("CALC%s:MARK7:Y?" % (i)))
            min_val4 = float(device.query("CALC%s:MARK8:Y?" % (i)))

            amp = (max([max_val1,max_val2,max_val3,max_val4]) - min([min_val1,min_val2,min_val3,min_val4]))/2
            amp_track.append(amp)

        for j in [2,4,6,8]:
            device.write(":CALCulate%s:PARameter:SELect 'Phase%s_Ch1'" % (j,j))
            max_val1 = float(device.query("CALC%s:MARK1:Y?" % (j)))
            min_val1 = float(device.query("CALC%s:MARK2:Y?" % (j)))

            device.write(":CALCulate%s:PARameter:SELect 'Phase%s_Ch2'" % (j,j))
            max_val2 = float(device.query("CALC%s:MARK3:Y?" % (j)))
            min_val2 = float(device.query("CALC%s:MARK4:Y?" % (j)))

            device.write(":CALCulate%s:PARameter:SELect 'Phase%s_Ch3'" % (j,j))
            max_val3 = float(device.query("CALC%s:MARK5:Y?" % (j)))
            min_val3 = float(device.query("CALC%s:MARK6:Y?" % (j)))

            device.write(":CALCulate%s:PARameter:SELect 'Phase%s_Ch4'" % (j,j))
            max_val4 = float(device.query("CALC%s:MARK7:Y?" % (j)))
            min_val4 = float(device.query("CALC%s:MARK8:Y?" % (j)))

            phase = (max([max_val1,max_val2,max_val3,max_val4]) - min([min_val1,min_val2,min_val3,min_val4]))/2
            phase_track.append(phase)

        self.export_sfb(amp_track,phase_track,band_num,bite)

        device.close()
        rm.close()
    
    def export_sfb(self,amp_track,phase_track,band_num=1,bite=False):
        now = datetime.now()
        dt_string = now.strftime("%d_%m_%Y_%H_%M")
        os.mkdir("test_data/%s" % dt_string)
        bands = [["0.5 - 0.8","0.7 - 1.2","1.1 - 1.6","1.5 - 2.2"],
                ["2.2 - 2.75","2.5 - 3.75","3.5 - 5.25","5 - 6.25"],
                ["6.0 - 9.5","9.0 - 13.0","12.5 - 14.5","14.0 - 18.0"]]

        sfb_dict = {"Frequency (GHz)":bands[band_num-1],
                    "Amplitude Tracking (Db)":[0.0,0.0,0.0,0.0],
                    "Phase Tracking (deg)":[0.0,0.0,0.0,0.0]
                    }

        sfb = pd.DataFrame(sfb_dict,index=[1,2,3,4])

        for a in range(4):
            sfb.iloc[a,1] = amp_track[a]
            sfb.iloc[a,2] = phase_track[a]
        if band_num == 1:
            dut = self.sfb1_sl
        elif band_num == 2:
            dut = self.sfb2_sl
        else:
            dut = self.sfb3_sl
        
        if bite:
            filename = "%s_SFB%s_BITE_%s" % (dt_string,band_num,dut)
        else:
            filename = "%s_SFB%s_RF_%s" % (dt_string,band_num,dut)
        
        ex = filename + ".xlsx"
        wo = filename + ".docx"
        pdf_name = filename + ".pdf"

        sfb.to_excel("test_data/%s/%s" % (dt_string,ex),index_label="Band")
        
        if bite:
            mode = "BITE"
        else:
            mode = "RF"

        template = "template_SFB%s_%s.docx" % (str(band_num),mode)

        self.report_sfb(dt_string,template,wo,pdf_name,sfb)

    def save_diagram(self,folder_name,bite,dut):
        now = datetime.now()
        dt_string = now.strftime("%d_%m_%Y_%H_%M")
        name_string = dt_string+"_"+folder_name+"_"+bite+"_"+dut+".znx"
        rm = visa.ResourceManager()
        device = rm.open_resource(self.ip_zna)

        device.write("MMEMory:CDIRectory 'C:\\Synergy'")
        device.write("MMEM:STOR:STAT 1,'%s\\%s\\%s'" % (folder_name,bite,name_string))

        device.close()
        rm.close()

    def data_trace_off(self):
        ip_zna = self.ip_zna

        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)

        for i in range(4):
            device.write("DISP:TRAC:SHOW 'Gain%s', OFF" % (2*i+1))
            time.sleep(0.5)
            device.write("DISP:TRAC:SHOW 'Phase%s', OFF" % (2*i+2))
            time.sleep(0.5)

        device.close()
        rm.close()

    ####QSRx Module####

    def submit_ps1_qsr(self):
        ip_ps = self.ip_ps

        rm = visa.ResourceManager()
        device = rm.open_resource(ip_ps)

        device.write('*RST')

        device.write("INST OUT1")
        time.sleep(0.1)
        device.write('OUTP:SEL 1')
        device.write('VOLT:PROT 15')
        device.write('VOLT 9')
        device.write('CURR 1')
        time.sleep(0.1)

        device.write("INST OUT2")
        time.sleep(0.1)
        device.write('OUTP:SEL 1')
        device.write('VOLT:PROT 15')
        device.write('VOLT 9')
        device.write('CURR 6')
        time.sleep(0.1)

        device.write("INST OUT3")
        time.sleep(0.1)
        device.write('OUTP:SEL 1')
        device.write('VOLT:PROT 20')
        device.write('VOLT 15')
        device.write('CURR 3')
        
        time.sleep(0.1)
        
        device.write("INST OUT4")
        time.sleep(0.1)
        device.write('OUTP:SEL 1')
        device.write('VOLT:PROT 7.5')
        device.write('VOLT 5')
        device.write('CURR 0.3')
        time.sleep(0.1)

        
        
        time.sleep(0.1)

        device.write('OUTP 1')
        
        device.close()
        rm.close()

        CTkMessagebox(title='Warning!', message='Please wait for 40 seconds before continuing.', option_1='OK')

    def mixer_settings_gain(self,bands,window):
        ip_zna = self.ip_zna
        
        start_f = bands[0]
        stop_f = bands[1]
        lo1 = bands[2]
        lo2 = bands[3]

        #self.lo_settings(lo1,14)

        offset = lo1-lo2

        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)

        #device.write('*RST')
        #device.write(':DISPlay:WINDow%s:STATe %d' % (window,1))
        
        time.sleep(0.5)
        #device.write("DISPlay:WINDow%s:STATe ON" % (window))

        device.write('SOUR:POWer1 -30 DBM')
        device.write("SOUR:POW1:ATT 20")
        device.write('SENS%s:SWEEp:POINts 101' % (window))
        time.sleep(0.5)
        device.write('SENS%s:FREQ:CONV:MIX:RFPort 1' % (window))
        device.write('SENS%s:FREQ:CONV:MIX:IFPort 2' % (window))
        device.write(':SENSe%s:BANDwidth:RESolution %fKHz' % (window,10))
        time.sleep(0.5)
        device.write('SOUR:FREQ:CONV:MIX:PMODE IF, FUNDamental')
        device.write('SENS%s:FREQ:CONV:MIX:STAG 2' % (window))
        device.write('SENS%s:FREQ:CONV:MIX:FIXED1 LO1' % (window))
        device.write('SENS%s:FREQ:CONV:MIX:FIXED2 LO2' % (window))
        device.write('SENS%s:FREQ:CONV:MIX:MFFixed LO1, %fGHz' % (window,lo1))
        device.write('SENS%s:FREQ:CONV:MIX:MFFixed LO2, %fGHz' % (window,lo2))
        time.sleep(0.5)
        device.write('SENS%s:FREQ:START %fGHz' % (window,start_f))
        device.write('SENS%s:FREQ:STOP %fGHz' % (window,stop_f))

        if start_f == 13.5 and lo1 == 16.75:
            device.write('SENS%s:FREQ:CONV:MIX:TFRequency1 DCL' % (window))
            device.write('SENS%s:FREQ:CONV:MIX:TFRequency2 DCL' % (window))
            offset = lo1 - lo2
            mul = 1
        
        elif start_f == 4.5 and lo1 == 16.75:
            device.write('SENS%s:FREQ:CONV:MIX:TFRequency1 DCL' % (window))
            device.write('SENS%s:FREQ:CONV:MIX:TFRequency2 DCUP' % (window))
            offset = lo2 - lo1
            mul = -1

        elif lo1 > start_f:
            mul = 1
            offset = lo1 - lo2
            device.write('SENS%s:FREQ:CONV:MIX:TFRequency1 DCL' % (window))
            device.write('SENS%s:FREQ:CONV:MIX:TFRequency2 DCL' % (window))
            time.sleep(0.5)
            device.write('SENS%s:FREQ:START %f' % (2,750e6))
            if start_f == 0.5:
                device.write('SENS%s:FREQ:STOP %f' % (2,1050e6))
            else:
                device.write('SENS%s:FREQ:STOP %f' % (2,1250e6))
        else:
            mul = -1
            offset = -(lo1 + lo2)
            device.write('SENS%s:FREQ:CONV:MIX:TFRequency1 DCUP' % (window))
            device.write('SENS%s:FREQ:CONV:MIX:TFRequency2 DCL' % (window))
            time.sleep(0.5)
            device.write('SENS%s:FREQ:START %f' % (2,750e6))
            device.write('SENS%s:FREQ:STOP %f' % (2,1250e6))
            

        #time.sleep(1)
        device.write('SENS%s:FREQuency:CONVersion MIXer' % (window))
        time.sleep(1)
        device.write("SENS%s:SWE:AXIS:FREQ 'Port 2; Receiver'" % (window))
        device.write("SOUR%s:FREQ4:CONV:ARB:IFR %d, 1, %fGHz , SWE" % (window,mul,-offset))
        
        device.write(":CONFigure:TRACe1:REName 'Trc%s', 'Gain%s'" % (window,window))
        device.write(":DISP:WIND%s:TITL:DATA 'Gain: %s - %sGHz'" % (window,start_f,stop_f))

        device.close()
        rm.close()

    def mixer_settings_phase(self,ref_port,bands,window):
        ip_zna = self.ip_zna

        start_f = bands[0]
        stop_f = bands[1]
        lo1 = bands[2]
        lo2 = bands[3]

        #self.lo_settings(lo1,14)

        offset = lo1-lo2

        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)

        #device.write('*RST')
        #device.write(':DISPlay:WINDow%s:STATe %d' % (window,1))
        
        time.sleep(0.5)
        #device.write("DISPlay:WINDow%s:STATe ON" % (window))

        device.write('SOUR:POWer1 -30 DBM')
        device.write("SOUR:POW1:ATT 20")
        device.write('SENS%s:SWEEp:POINts 101' % (window))
        time.sleep(0.5)
        device.write('SENS%s:FREQ:CONV:MIX:RFPort 1' % (window))
        device.write('SENS%s:FREQ:CONV:MIX:IFPort 2' % (window))
        device.write(':SENSe%s:BANDwidth:RESolution %fKHz' % (window,10))
        time.sleep(0.5)
        device.write('SOUR:FREQ:CONV:MIX:PMODE IF, FUNDamental')
        device.write('SENS%s:FREQ:CONV:MIX:STAG 2' % (window))
        device.write('SENS%s:FREQ:CONV:MIX:FIXED1 LO1' % (window))
        device.write('SENS%s:FREQ:CONV:MIX:FIXED2 LO2' % (window))
        device.write('SENS%s:FREQ:CONV:MIX:MFFixed LO1, %fGHz' % (window,lo1))
        device.write('SENS%s:FREQ:CONV:MIX:MFFixed LO2, %fGHz' % (window,lo2))
        time.sleep(0.5)
        device.write('SENS%s:FREQ:START %fGHz' % (window,start_f))
        device.write('SENS%s:FREQ:STOP %fGHz' % (window,stop_f))

        if start_f == 13.5 and lo1 == 16.75:
            device.write('SENS%s:FREQ:CONV:MIX:TFRequency1 DCL' % (window))
            device.write('SENS%s:FREQ:CONV:MIX:TFRequency2 DCL' % (window))
            offset = lo1 - lo2
            mul = 1
        
        elif start_f == 4.5 and lo1 == 16.75:
            device.write('SENS%s:FREQ:CONV:MIX:TFRequency1 DCL' % (window))
            device.write('SENS%s:FREQ:CONV:MIX:TFRequency2 DCUP' % (window))
            offset = lo2 - lo1
            mul = -1

        elif lo1 > start_f:
            mul = 1
            offset = lo1 - lo2
            device.write('SENS%s:FREQ:CONV:MIX:TFRequency1 DCL' % (window))
            device.write('SENS%s:FREQ:CONV:MIX:TFRequency2 DCL' % (window))
            time.sleep(0.5)
            device.write('SENS%s:FREQ:START %f' % (2,750e6))
            if start_f == 0.5:
                device.write('SENS%s:FREQ:STOP %f' % (2,1050e6))
            else:
                device.write('SENS%s:FREQ:STOP %f' % (2,1250e6))
        elif lo1 < start_f:
            mul = -1
            offset = -(lo1 + lo2)
            device.write('SENS%s:FREQ:CONV:MIX:TFRequency1 DCUP' % (window))
            device.write('SENS%s:FREQ:CONV:MIX:TFRequency2 DCL' % (window))
            time.sleep(0.5)
            device.write('SENS%s:FREQ:START %f' % (2,1250e6))
            device.write('SENS%s:FREQ:STOP %f' % (2,750e6))
            

        #time.sleep(1)
        device.write('SENS%s:FREQuency:CONVersion MIXer' % (window))
        time.sleep(1)
        device.write("SENS%s:SWE:AXIS:FREQ 'Port 2; Receiver'" % (window))
        device.write("SOUR%s:FREQ4:CONV:ARB:IFR %d, 1, %fGHz , SWE" % (window,mul,-offset))
        
        
        device.write(":CONFigure:TRACe1:REName 'Trc%s', 'Gain%s'" % (window,window))
        
        device.write(":DISP:WIND%s:TITL:DATA 'Phase(Ref.%s): %s - %sGHz'" % (window,ref_port,start_f,stop_f))

        device.close()
        rm.close()

    def phase_def(self,window,ref_port):
        ip_zna = self.ip_zna
        
        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)

        device.write(":CALC%s:PAR:SDEF 'Phase%s_R%s', 'B2/B4'" % (window,window,ref_port))
        device.write(":DISP:WIND%s:TRAC:EFEED 'Phase%s_R%s'" % (window,window,ref_port))
        time.sleep(1)
        device.write(":CALC%s:PAR:DEL 'Gain%s'" % (window,window))
        device.write(':CALCulate%s:FORMat %s' % (window,'PHAse'))
        
        device.close()
        rm.close()

    def phase_ref_qsrx(self,lo1,ref_port,window):
        ip_zna = self.ip_zna

        self.lo_settings(lo1,14)
        #mixer_settings(band,window)    
        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)
        
        #device.write(":DISPlay:WINDow%s:TRACe:Y:PDIVision 5,'Phase%s_R%s'" % (window,window,ref_port))
        #device.write("DISP:WIND%s:TRAC:Y:RLEV 0, 'Phase%s_R%s'" % (window,window,ref_port))
        #device.write("DISP:WIND%s:TRAC:Y:RPOS 50, 'Phase%s_R%s'" % (window,window,ref_port))
        
        time.sleep(3)
        device.write(":TRAC:COPY:MATH 'Phase%s_R%s_mem','Phase%s_R%s'" % (window,ref_port,window,ref_port))
        time.sleep(1)
        device.write(":DISP:WIND%s:TRAC:EFEED 'Phase%s_R%s_mem'" % (window,window,ref_port))
        time.sleep(1)
        
        device.write(":CALC%s:MATH:SDEF 'Phase%s_R%s / Phase%s_R%s_mem'" % (window,window,ref_port,window,ref_port))
        time.sleep(1)
        device.write(':CALC%s:MATH:STAT ON' % (window))
        time.sleep(1)
        device.write("TRAC:COPY:MATH 'Phase%s_R%s_Ch1','Phase%s_R%s'" % (window,ref_port,window,ref_port))
        time.sleep(1)
        #device.write(":CALC:MATH:STAT OFF")
        device.write(":DISP:WIND%s:TRAC:EFEED 'Phase%s_R%s_Ch1'" % (window,window,ref_port))
        time.sleep(0.5)
        device.write(":DISPlay:WINDow%s:TRACe:Y:PDIVision 7,'Phase%s_R%s_Ch1'" % (window,window,ref_port))
        device.write("DISP:WIND%s:TRAC:Y:RLEV 0, 'Phase%s_R%s_Ch1'" % (window,window,ref_port))
        device.write("DISP:WIND%s:TRAC:Y:RPOS 50, 'Phase%s_R%s_Ch1'" % (window,window,ref_port))
        
        #device.write("DISP:TRAC:SHOW 'Phase%s_R%s_mem', OFF" % (window,ref_port))
        
        device.close()
        rm.close()

    def gain_ref_qsrx(self,lo1,window):
        ip_zna = self.ip_zna

        self.lo_settings(lo1,14)
        #mixer_settings(band,window)

        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)
        
        device.write(':CALCulate%s:FORMat %s' % (window,'MLOGarithmic'))
        time.sleep(0.5)
        device.write(":TRAC:COPY:MATH 'Gain%s_mem','Gain%s'" % (window,window))
        time.sleep(0.5)
        device.write(":DISP:WIND%s:TRAC:EFEED 'Gain%s_mem'" % (window,window))
        time.sleep(0.5)
        device.write(":CALC%s:MATH:SDEF 'Gain%s / Gain%s_mem'" % (window,window,window))
        device.write(':CALC%s:MATH:STAT ON' % (window))
        time.sleep(0.5)
        device.write("TRAC:COPY:MATH 'Gain%s_Ch1','Gain%s'" % (window,window))
        time.sleep(0.5)
        #device.write(":CALC:MATH:STAT OFF")
        device.write(":DISP:WIND%s:TRAC:EFEED 'Gain%s_Ch1'" % (window,window))
        device.write(":DISPlay:WINDow%s:TRACe:Y:PDIVision 2,'Gain%s_Ch1'" % (window,window))
        device.write("DISP:WIND%s:TRAC:Y:RLEV 0, 'Gain%s_Ch1'" % (window,window))
        device.write("DISP:WIND%s:TRAC:Y:RPOS 50, 'Gain%s_Ch1'" % (window,window))
        #device.write("DISP:TRAC:SHOW 'Gain%s_mem', OFF" % (window))

        device.close()
        rm.close()

    def new_phase_qsrx(self,lo1,window,ref_port,p_num):
        ip_zna = self.ip_zna

        self.lo_settings(lo1,14)

        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)

        device.write(":CALCulate%s:PARameter:SELect 'Phase%s_R%s'" % (window,window,ref_port))
        time.sleep(0.5)
        device.write(":CALC%s:MATH:SDEF 'Phase%s_R%s / Phase%s_R%s_mem'" % (window,window,ref_port,window,ref_port))
        time.sleep(0.5)
        device.write(':CALC%s:MATH:STAT ON' % (window))
        device.write("TRAC:COPY:MATH 'Phase%s_R%s_Ch%s','Phase%s_R%s'" % (window,ref_port,p_num,window,ref_port))
        time.sleep(0.5)
        #device.write(":CALC:MATH:STAT OFF")
        #device.write("DISPlay:WINDow2:STATe ON")
        device.write(":DISP:WIND%s:TRAC:EFEED 'Phase%s_R%s_Ch%s'" % (window,window,ref_port,p_num))
        time.sleep(0.5)
        device.write(":DISPlay:WINDow%s:TRACe:Y:PDIVision 7,'Phase%s_R%s_Ch%s'" % (window,window,ref_port,p_num))
        device.write("DISP:WIND%s:TRAC:Y:RLEV 0, 'Phase%s_R%s_Ch%s'" % (window,window,ref_port,p_num))
        device.write("DISP:WIND%s:TRAC:Y:RPOS 50, 'Phase%s_R%s_Ch%s'" % (window,window,ref_port,p_num))
        time.sleep(1)

        device.close()
        rm.close()

    def lo_settings(self,freq,amp):
        time.sleep(1)
        ip_lo = self.ip_lo

        rm = visa.ResourceManager()
        device_lo = rm.open_resource(ip_lo)
        device_lo.write(':SOURce:FREQuency:FIXed %fGhz' % (freq))
        device_lo.write(':SOURce:POWer:LEVel:IMMediate:AMPLitude %G' % (amp))
        device_lo.write(':OUTPut:MODulation:STATe %d' % (0))
        device_lo.write(':OUTPut:STATe %d' % (1))
        device_lo.close()
        rm.close()

        time.sleep(3)

    def new_gain_qsrx(self,lo1,window,g_num):
        ip_zna = self.ip_zna

        self.lo_settings(lo1,14)
        #time.sleep(2)
        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)

        
        device.write(":CALCulate%s:PARameter:SELect 'Gain%s'" % (window,window))
        device.write(":CALC%s:MATH:SDEF 'Gain%s / Gain%s_mem'" % (window,window,window))
        time.sleep(0.5)
        device.write(':CALC%s:MATH:STAT ON' % (window))
        time.sleep(0.5)
        device.write("TRAC:COPY:MATH 'Gain%s_Ch%s','Gain%s'" % (window,g_num,window))
        time.sleep(0.5)
        #device.write(":CALC:MATH:STAT OFF")
        #device.write("DISPlay:WINDow2:STATe ON")
        device.write(":DISP:WIND%s:TRAC:EFEED 'Gain%s_Ch%s'" % (window,window,g_num))
        time.sleep(0.5)
        device.write(":DISPlay:WINDow%s:TRACe:Y:PDIVision 2,'Gain%s_Ch%s'" % (window,window,g_num))
        device.write("DISP:WIND%s:TRAC:Y:RLEV 0, 'Gain%s_Ch%s'" % (window,window,g_num))
        device.write("DISP:WIND%s:TRAC:Y:RPOS 50, 'Gain%s_Ch%s'" % (window,window,g_num))
        time.sleep(0.5)
        
        device.close()
        rm.close()

    def marker_qsr(self,num):
        ip_zna = self.ip_zna

        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)

        #device.write('*RST')
        amp_track = []
        phase_track_r2 = []
        phase_track_r4 = []

        for i in range(1,num*3+1,3):
            for j in range(4):
                device.write(":CALCulate%s:PARameter:SELect 'Gain%s_Ch%s'" % (i,i,j+1))
                device.write(":CALCulate%s:MARKer%s ON" % (i,2*j+1))
                device.write(":CALCulate%s:MARKer%s:MAX" % (i,2*j+1))

                device.write(":CALCulate%s:MARKer%s ON" % (i,2*j+2))
                device.write(":CALCulate%s:MARKer%s:MIN" % (i,2*j+2))
                time.sleep(0.5)

        for m in range(2,num*3+1,3):
            
            device.write(":CALCulate%s:PARameter:SELect 'Phase%s_R2_Ch1'" % (m,m))
            device.write(":CALCulate%s:MARKer%s ON" % (m,1))
            device.write(":CALCulate%s:MARKer%s:MAX" % (m,1))

            device.write(":CALCulate%s:MARKer%s ON" % (m,2))
            device.write(":CALCulate%s:MARKer%s:MIN" % (m,2))

            time.sleep(0.5)

            device.write(":CALCulate%s:PARameter:SELect 'Phase%s_R2_Ch3'" % (m,m))
            device.write(":CALCulate%s:MARKer%s ON" % (m,3))
            device.write(":CALCulate%s:MARKer%s:MAX" % (m,3))

            device.write(":CALCulate%s:MARKer%s ON" % (m,4))
            device.write(":CALCulate%s:MARKer%s:MIN" % (m,4))

            time.sleep(0.5)

            device.write(":CALCulate%s:PARameter:SELect 'Phase%s_R2_Ch4'" % (m,m))
            device.write(":CALCulate%s:MARKer%s ON" % (m,5))
            device.write(":CALCulate%s:MARKer%s:MAX" % (m,5))

            device.write(":CALCulate%s:MARKer%s ON" % (m,6))
            device.write(":CALCulate%s:MARKer%s:MIN" % (m,6))

            time.sleep(0.5)

        for n in range(3,num*3+1,3):
            
            device.write(":CALCulate%s:PARameter:SELect 'Phase%s_R4_Ch1'" % (n,n))
            device.write(":CALCulate%s:MARKer%s ON" % (n,1))
            device.write(":CALCulate%s:MARKer%s:MAX" % (n,1))

            device.write(":CALCulate%s:MARKer%s ON" % (n,2))
            device.write(":CALCulate%s:MARKer%s:MIN" % (n,2))

            time.sleep(0.5)

            device.write(":CALCulate%s:PARameter:SELect 'Phase%s_R4_Ch2'" % (n,n))
            device.write(":CALCulate%s:MARKer%s ON" % (n,3))
            device.write(":CALCulate%s:MARKer%s:MAX" % (n,3))

            device.write(":CALCulate%s:MARKer%s ON" % (n,4))
            device.write(":CALCulate%s:MARKer%s:MIN" % (n,4))

            time.sleep(0.5)

            device.write(":CALCulate%s:PARameter:SELect 'Phase%s_R4_Ch3'" % (n,n))
            device.write(":CALCulate%s:MARKer%s ON" % (n,5))
            device.write(":CALCulate%s:MARKer%s:MAX" % (n,5))

            device.write(":CALCulate%s:MARKer%s ON" % (n,6))
            device.write(":CALCulate%s:MARKer%s:MIN" % (n,6))

            time.sleep(0.5)
        
        for a1 in range(1,num*3+1,3):
            max_val1 = []
            min_val1 = []
            for a2 in range(4):
                device.write(":CALCulate%s:PARameter:SELect 'Gain%s_Ch%s'" % (a1,a1,a2+1))
                max_val1.append(float(device.query("CALC%s:MARK%s:Y?" % (a1,2*a2+1))))
                min_val1.append(float(device.query("CALC%s:MARK%s:Y?" % (a1,2*a2+2))))

            amp = (max(max_val1) - min(min_val1))/2
            amp_track.append(amp)

        for b1 in range(2,num*3+1,3):
            max_val2 = []
            min_val2 = []
            
            device.write(":CALCulate%s:PARameter:SELect 'Phase%s_R2_Ch1'" % (b1,b1))
            max_val2.append(float(device.query("CALC%s:MARK%s:Y?" % (b1,1))))
            min_val2.append(float(device.query("CALC%s:MARK%s:Y?" % (b1,2))))

            device.write(":CALCulate%s:PARameter:SELect 'Phase%s_R2_Ch3'" % (b1,b1))
            max_val2.append(float(device.query("CALC%s:MARK%s:Y?" % (b1,3))))
            min_val2.append(float(device.query("CALC%s:MARK%s:Y?" % (b1,4))))

            device.write(":CALCulate%s:PARameter:SELect 'Phase%s_R2_Ch4'" % (b1,b1))
            max_val2.append(float(device.query("CALC%s:MARK%s:Y?" % (b1,5))))
            min_val2.append(float(device.query("CALC%s:MARK%s:Y?" % (b1,6))))

            p1 = (max(max_val2) - min(min_val2))/2
            phase_track_r2.append(p1)

        for c1 in range(3,num*3+1,3):
            max_val3 = []
            min_val3 = []
            
            device.write(":CALCulate%s:PARameter:SELect 'Phase%s_R4_Ch1'" % (c1,c1))
            max_val3.append(float(device.query("CALC%s:MARK%s:Y?" % (c1,1))))
            min_val3.append(float(device.query("CALC%s:MARK%s:Y?" % (c1,2))))

            device.write(":CALCulate%s:PARameter:SELect 'Phase%s_R4_Ch2'" % (c1,c1))
            max_val3.append(float(device.query("CALC%s:MARK%s:Y?" % (c1,3))))
            min_val3.append(float(device.query("CALC%s:MARK%s:Y?" % (c1,4))))

            device.write(":CALCulate%s:PARameter:SELect 'Phase%s_R4_Ch3'" % (c1,c1))
            max_val3.append(float(device.query("CALC%s:MARK%s:Y?" % (c1,5))))
            min_val3.append(float(device.query("CALC%s:MARK%s:Y?" % (c1,6))))

            p2 = (max(max_val3) - min(min_val3))/2
            phase_track_r4.append(p2)

        return amp_track,phase_track_r2,phase_track_r4

    def new_folder(self,dut,integrated=False,bite=False):
        ip_zna = self.ip_zna

        now = datetime.now()
        dt_string = now.strftime("%d_%m_%Y_%H_%M")
        name_string = dt_string +"_"+dut
        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)

        if integrated == False:
            folder_path = 'C:\\Synergy\\QSRx'
        elif bite == False:
            folder_path = 'C:\\Synergy\\QSRx_int\\RF'
        else:
            folder_path = 'C:\\Synergy\\QSRx_int\\BITE'
            
        
        device.write("MMEM:MDIR '%s\\%s'" % (folder_path,name_string))

        
        device.close()
        rm.close()

        return name_string

    def qsrx_band1(self,dt_string):
        ip_zna = self.ip_zna

        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()

        qsrx = [[0.5,0.8,14.75,15],
                [0.7,1.2,14.55,14.6],
                [1.1,1.6,14.15,13.8],
                [1.5,2,13.75,13]]
        
        controls = ["1 1","1 0","0 1","0 0"]

        current_band = qsrx

        ref_1 = self.mbox("Current Band is %s GHz - %s GHz\nConnect VNA Port-1 to CH-1 Input of Band 1\nConnect VNA Port-2 to CH-1 DRx Output\nConnect VNA Ref Port to CH-2 Input of Band 1\nConnect VNA Port-4 to CH-2 DRx Output\nSet QSRx Band-1 Control : 1 1" % (0.5,2.2))
        
        if ref_1 == 'OK':
            #Window 1
            for step1 in range(4):
                if self.stop_event.is_set():
                    return
                
                self.mixer_settings_gain(current_band[step1],3*step1+1)
                self.mixer_settings_phase(2,current_band[step1],3*step1+2)
                self.mixer_settings_phase(4,current_band[step1],3*step1+3)
                
            time.sleep(2)

            rm = visa.ResourceManager()
            device = rm.open_resource(ip_zna)
            device.write('SENS%s:FREQ:START %f' % (2,750e6))
            device.write('SENS%s:FREQ:STOP %f' % (2,1050e6))
            device.write('SENS%s:FREQ:START %f' % (3,750e6))
            device.write('SENS%s:FREQ:STOP %f' % (3,1050e6))
            device.close()
            rm.close()

            for step2 in range(4):
                if self.stop_event.is_set():
                    return
                win_num = 3*step2+2
                self.phase_def(win_num,2)

            for step3 in range(4):
                if self.stop_event.is_set():
                    return
                win_num = 3*step3+3
                self.phase_def(win_num,4)

            for step4 in range(4):
                if self.stop_event.is_set():
                    return
                pop = 'OK'
                if pop == 'OK':
                    win_num1 = 3*step4+2
                    win_num2 = 3*step4+1
                    self.phase_ref_qsrx(current_band[step4][2],2,win_num1)
                    self.gain_ref_qsrx(current_band[step4][2],win_num2)
            
            rm = visa.ResourceManager()
            device = rm.open_resource(ip_zna)
            device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx\\%s\\0.5-2.2GHz.znx'" % (dt_string))
            time.sleep(2)
            #device.write(':SYSTem:PRESet')
            device.close()
            rm.close()

            for chan in [3,4]:
                if self.stop_event.is_set():
                    return

                chan_change = self.mbox("Connect VNA Port-1 to CH-%s DRx Input\nConnect VNA Port-2 to CH-%s DRx Output" % (chan,chan))
                if chan_change == 'OK':
                    #Window 1
                    
                    for step5 in range(4):
                        if self.stop_event.is_set():
                            return
                        pop = 'OK'
                        if pop == 'OK':
                            win1 = 3*step5+1
                            self.new_gain_qsrx(current_band[step5][2],win1,chan)
                            win2 = 3*step5+2
                            self.new_phase_qsrx(current_band[step5][2],win2,2,chan)
                else:
                    self.stop_task()


            ref_change = self.mbox("Connect VNA Port-1 to CH-1 Input of Band 1\nConnect VNA Port-2 to CH-1 DRx Output\nConnect VNA Ref Port to CH-4 DRx Input\nConnect VNA Port-4 to CH-4 DRx Output")

            if ref_change == 'OK':

                for step6 in range(4):
                    if self.stop_event.is_set():
                        return
                    pop = 'OK'
                    if pop == 'OK':
                        win_num = 3*step6+3
                        self.phase_ref_qsrx(current_band[step6][2],4,win_num)


                for chan in [2,3]:
                    if self.stop_event.is_set():
                        return
                    chan_change = self.mbox("Connect VNA Port-1 to CH-%s DRx Input\nConnect VNA Port-2 to CH-%s DRx Output" % (chan,chan))
                    if chan_change == 'OK':

                        for step7 in range(4):
                            if self.stop_event.is_set():
                                return
                            pop = 'OK'
                            if pop == 'OK':
                                win1 = 3*step7+1
                                if chan == 2:
                                    self.new_gain_qsrx(current_band[step7][2],win1,chan)
                                win2 = 3*step7+3
                                self.new_phase_qsrx(current_band[step7][2],win2,4,chan)

                    else:
                        self.stop_task()
                        break
                        
                x,y,z = self.marker_qsr(4)
                rm = visa.ResourceManager()
                device = rm.open_resource(ip_zna)
                device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx\\%s\\0.5-2.2GHz.znx'" % (dt_string))
                time.sleep(2)
                device.close()
                rm.close()

                return x,y,z
            
            else:
                self.stop_task()

        else:
            self.stop_task()

    def qsrx_band2(self,dt_string):
        ip_zna = self.ip_zna

        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()

        qsrx = [[2,2.5,13.25,12],
                [2.5,3,12.75,11],
                [3,3.5,12.25,10],
                [3.5,4,11.75,9],
                [4,4.5,11.25,8],
                [4.5,5,10.75,7],
                [4.5,5,16.75,11],
                [5,5.5,10.25,6],
                [5.5,6,9.75,5]]
        
        controls = ["0 1","1 0","0 0","1 1"]

        current_band = qsrx

        ref_1 = self.mbox("Current Band is %s GHz - %s GHz\nConnect VNA Port-1 to CH-1 Input of Band 2\nConnect VNA Port-2 to CH-1 DRx Output\nConnect VNA Ref Port to CH-2 Input of Band 2\nConnect VNA Port-4 to CH-2 DRx Output\nSet QSRx Band-2 Control : 0 1" % (2,6.25))
        
        if ref_1 == 'OK':
            #Window 1
            for step1 in range(4):
                if self.stop_event.is_set():
                    return
                self.mixer_settings_gain(current_band[step1],3*step1+1)
                self.mixer_settings_phase(2,current_band[step1],3*step1+2)
                self.mixer_settings_phase(4,current_band[step1],3*step1+3)

            time.sleep(2)

            rm = visa.ResourceManager()
            device = rm.open_resource(ip_zna)
            device.write('SENS%s:FREQ:START %f' % (3,750e6))
            device.write('SENS%s:FREQ:STOP %f' % (3,1250e6))
            device.close()
            rm.close()

            for step2 in range(4):
                if self.stop_event.is_set():
                    return
                win_num = 3*step2+2
                self.phase_def(win_num,2)

            for step3 in range(4):
                if self.stop_event.is_set():
                    return
                win_num = 3*step3+3
                self.phase_def(win_num,4)

            for step4 in range(4):
                if self.stop_event.is_set():
                    return
                win_num1 = 3*step4+2
                win_num2 = 3*step4+1
                
                pop4 = 'OK'
                
                if pop4 == 'OK':
                    self.phase_ref_qsrx(current_band[step4][2],2,win_num1)
                    time.sleep(2)
                    self.gain_ref_qsrx(current_band[step4][2],win_num2)
                else:
                    self.stop_task()

            rm = visa.ResourceManager()
            device = rm.open_resource(ip_zna)
            device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx\\%s\\2-4GHz.znx'" % (dt_string))
            time.sleep(2)
            #device.write(':SYSTem:PRESet')
            device.close()
            rm.close()
            
            #Window 2
            rm = visa.ResourceManager()
            device = rm.open_resource(ip_zna)
            device.write(':SYSTem:PRESet')
            device.close()
            rm.close()

            for step1 in range(5):
                if self.stop_event.is_set():
                    return
                self.mixer_settings_gain(current_band[step1+4],3*step1+1)
                self.mixer_settings_phase(2,current_band[step1+4],3*step1+2)
                self.mixer_settings_phase(4,current_band[step1+4],3*step1+3)
                
            time.sleep(5)

            rm = visa.ResourceManager()
            device = rm.open_resource(ip_zna)
            device.write('SENS%s:FREQ:START %f' % (3,750e6))
            device.write('SENS%s:FREQ:STOP %f' % (3,1250e6))
            device.close()
            rm.close()

            for step2 in range(5):
                if self.stop_event.is_set():
                    return
                win_num = 3*step2+2
                self.phase_def(win_num,2)

            for step3 in range(5):
                if self.stop_event.is_set():
                    return
                win_num = 3*step3+3
                self.phase_def(win_num,4)

            for step4 in range(5):
                if self.stop_event.is_set():
                    return
                win_num1 = 3*step4+2
                win_num2 = 3*step4+1
                pop4 = 'OK'

                if pop4 == 'OK':
                    self.phase_ref_qsrx(current_band[step4+4][2],2,win_num1)
                    time.sleep(2)
                    self.gain_ref_qsrx(current_band[step4+4][2],win_num2)
                else:
                    self.stop_task()

            rm = visa.ResourceManager()
            device = rm.open_resource(ip_zna)
            device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx\\%s\\4-6GHz.znx'" % (dt_string))
            time.sleep(1)
            device.close()
            rm.close()

            rm = visa.ResourceManager()
            device = rm.open_resource(ip_zna)
            #device.write(':SYSTem:PRESet')
            device.close()
            rm.close()

            for chan in [3,4]:
                if self.stop_event.is_set():
                    return

                chan_change = self.mbox("Connect VNA Port-1 to CH-%s Input of Band 2\nConnect VNA Port-2 to CH-%s DRx Output" % (chan,chan))
                if chan_change == 'OK':
                    #Window 1
                    rm = visa.ResourceManager()
                    device = rm.open_resource(ip_zna)
                    device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx\\%s\\2-4GHz.znx'" % (dt_string))
                    time.sleep(5)
                    device.close()
                    rm.close()
                    for step6 in range(4):
                        if self.stop_event.is_set():
                            return
                        
                        pop6 = 'OK'
                        
                        if pop6 == 'OK':
                            win1 = 3*step6+1
                            self.new_gain_qsrx(current_band[step6][2],win1,chan)
                            win2 = 3*step6+2
                            self.new_phase_qsrx(current_band[step6][2],win2,2,chan)
                        else:
                            self.stop_task()


                    #Window 2
                    rm = visa.ResourceManager()
                    device = rm.open_resource(ip_zna)
                    device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx\\%s\\2-4GHz.znx'" % (dt_string))
                    time.sleep(2)
                    device.close()
                    rm.close()

                    rm = visa.ResourceManager()
                    device = rm.open_resource(ip_zna)
                    #device.write(':SYSTem:PRESet')
                    time.sleep(2)
                    device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx\\%s\\4-6GHz.znx'" % (dt_string))
                    time.sleep(5)
                    device.close()
                    rm.close()
                    for step6 in range(5):
                        if self.stop_event.is_set():
                            return
                        win1 = 3*step6+1
                        
                        pop6 = 'OK'

                        if pop6 == 'OK':
                            self.new_gain_qsrx(current_band[step6+4][2],win1,chan)
                            win2 = 3*step6+2
                            self.new_phase_qsrx(current_band[step6+4][2],win2,2,chan)
                        else:
                            self.stop_task()

                    rm = visa.ResourceManager()
                    device = rm.open_resource(ip_zna)
                    device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx\\%s\\4-6GHz.znx'" % (dt_string))
                    time.sleep(2)
                    #device.write(':SYSTem:PRESet')
                    device.close()
                    rm.close()

                    rm = visa.ResourceManager()
                    device = rm.open_resource(ip_zna)
                    #device.write(':SYSTem:PRESet')
                    device.close()
                    rm.close()
                
                else:
                    self.stop_task()
                    break
                



            ref_change = self.mbox("Connect VNA Port-1 to CH-1 Input of Band 2\nConnect VNA Port-2 to CH-1 DRx Output\nConnect VNA Ref Port to CH-4 Input of Band 2\nConnect VNA Port-4 to CH-4 DRx Output")

            if ref_change == 'OK':
                rm = visa.ResourceManager()
                device = rm.open_resource(ip_zna)
                device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx\\%s\\2-4GHz.znx'" % (dt_string))
                time.sleep(5)
                device.close()
                rm.close()

                for step7 in range(4):
                    if self.stop_event.is_set():
                        return
                    win_num = 3*step7+3
                    
                    pop7 = 'OK'

                    if pop7 == 'OK':
                        self.phase_ref_qsrx(current_band[step7][2],4,win_num)

                rm = visa.ResourceManager()
                device = rm.open_resource(ip_zna)
                device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx\\%s\\2-4GHz.znx'" % (dt_string))
                time.sleep(2)
                device.close()
                rm.close()

                rm = visa.ResourceManager()
                device = rm.open_resource(ip_zna)
                #device.write(':SYSTem:PRESet')
                time.sleep(2)
                device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx\\%s\\4-6GHz.znx'" % (dt_string))
                time.sleep(5)
                device.close()
                rm.close()

                for step7 in range(5):
                    if self.stop_event.is_set():
                        return
                    
                    pop7 = 'OK'
                    
                    if pop7 == 'OK':
                        win_num = 3*step7+3
                        self.phase_ref_qsrx(current_band[step7+4][2],4,win_num)

                rm = visa.ResourceManager()
                device = rm.open_resource(ip_zna)
                device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx\\%s\\4-6GHz.znx'" % (dt_string))
                time.sleep(2)
                device.close()
                rm.close()

                for chan in [2,3]:
                    if self.stop_event.is_set():
                        return
                    chan_change = self.mbox("Connect VNA Port-1 to CH-%s Input of Band 2\nConnect VNA Port-2 to CH-%s DRx Output" % (chan,chan))
                    if chan_change == 'OK':
                        #Window 1
                        rm = visa.ResourceManager()
                        device = rm.open_resource(ip_zna)
                        device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx\\%s\\2-4GHz.znx'" % (dt_string))
                        time.sleep(5)
                        device.close()
                        rm.close()
                        for step8 in range(4):
                            win1 = 3*step8+1
                            
                            pop8 = 'OK'
                                
                            if pop8 == 'OK':
                                if chan == 2:
                                    self.new_gain_qsrx(current_band[step8][2],win1,chan)
                                win2 = 3*step8+3
                                self.new_phase_qsrx(current_band[step8][2],win2,4,chan)


                        #Window 2
                        rm = visa.ResourceManager()
                        device = rm.open_resource(ip_zna)
                        device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx\\%s\\2-4GHz.znx'" % (dt_string))
                        time.sleep(2)
                        device.close()
                        rm.close()

                        rm = visa.ResourceManager()
                        device = rm.open_resource(ip_zna)
                        #device.write(':SYSTem:PRESet')
                        time.sleep(2)
                        device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx\\%s\\4-6GHz.znx'" % (dt_string))
                        time.sleep(5)
                        device.close()
                        rm.close()
                        for step8 in range(5):
                            if self.stop_event.is_set():
                                return
                            win1 = 3*step8+1
                            
                            pop8 = 'OK'

                            if pop8 == 'OK':
                                if chan == 2:
                                    self.new_gain_qsrx(current_band[step8+4][2],win1,chan)
                                win2 = 3*step8+3
                                self.new_phase_qsrx(current_band[step8+4][2],win2,4,chan)

                        rm = visa.ResourceManager()
                        device = rm.open_resource(ip_zna)
                        device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx\\%s\\4-6GHz.znx'" % (dt_string))
                        time.sleep(2)
                        #device.write(':SYSTem:PRESet')
                        device.close()
                        rm.close()
                    else:
                        break
            
            rm = visa.ResourceManager()
            device = rm.open_resource(ip_zna)
            device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx\\%s\\2-4GHz.znx'" % (dt_string))
            time.sleep(2)
            #device.write(':SYSTem:PRESet')
            device.close()
            rm.close()

            x1,y1,z1 = self.marker_qsr(4)

            rm = visa.ResourceManager()
            device = rm.open_resource(ip_zna)
            device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx\\%s\\2-4GHz.znx'" % (dt_string))
            time.sleep(2)
            #device.write(':SYSTem:PRESet')
            device.close()
            rm.close()

            rm = visa.ResourceManager()
            device = rm.open_resource(ip_zna)
            device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx\\%s\\4-6GHz.znx'" % (dt_string))
            time.sleep(2)
            #device.write(':SYSTem:PRESet')
            device.close()
            rm.close()

            x2,y2,z2 = self.marker_qsr(5)

            rm = visa.ResourceManager()
            device = rm.open_resource(ip_zna)
            device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx\\%s\\4-6GHz.znx'" % (dt_string))
            time.sleep(2)
            #device.write(':SYSTem:PRESet')
            device.close()
            rm.close()

            x1.extend(x2)
            y1.extend(y2)
            z1.extend(z2)

            return x1,y1,z1
        
        else:
            self.stop_task()

    def qsrx_band3(self,dt_string):
        ip_zna = self.ip_zna

        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()

        qsrx = [[6,6.5,9.25,4],
                [6.5,7,9.75,4],
                [7,7.5,10.25,4],
                [7.5,8,10.75,4],
                [8,8.5,11.25,4],
                [8.5,9,11.75,4],
                [9,9.5,12.25,4],
                [9.5,10,12.75,4],
                [10,10.5,13.25,4],
                [10.5,11,13.75,4],
                [11,11.5,14.25,4],
                [11.5,12,14.75,4],
                [12,12.5,9.25,4],
                [12.5,13,9.75,4],
                [13,13.5,10.25,4],
                [13.5,14,10.75,4],
                [13.5,14,16.75,4],
                [14,14.5,11.25,4],
                [14.5,15,11.75,4],
                [15,15.5,12.25,4],
                [15.5,16,12.75,4],
                [16,16.5,13.25,4],
                [16.5,17,13.75,4],
                [17,17.5,14.25,4],
                [17.5,18,14.75,4]]
        
        controls = ["0 0","0 1","1 0","1 1"]

        current_band = qsrx

        bands = [[6,8.5],
                [8.5,11],
                [11,13.5],
                [13.5,15.5],
                [15.5,18]]

        ref_1 = self.mbox("Current Band is %s GHz - %s GHz\nConnect VNA Port-1 to CH-1 Input of Band 3\nConnect VNA Port-2 to CH-1 DRx Output\nConnect VNA Ref Port to CH-2 Input of Band 3\nConnect VNA Port-4 to CH-2 DRx Output\nSet QSRx Band-2 Control : 0 0" % (6,18))
        

        if ref_1 == 'OK':
            #Window 1
            for i in range(5):
                if self.stop_event.is_set():
                    return
                rm = visa.ResourceManager()
                device = rm.open_resource(ip_zna)
                device.write(':SYSTem:PRESet')
                device.close()
                rm.close()

                for step1 in range(5):
                    if self.stop_event.is_set():
                        return
                    self.mixer_settings_gain(current_band[step1+i*5],3*step1+1)
                    self.mixer_settings_phase(2,current_band[step1+i*5],3*step1+2)
                    self.mixer_settings_phase(4,current_band[step1+i*5],3*step1+3)
                    
                time.sleep(2)

                rm = visa.ResourceManager()
                device = rm.open_resource(ip_zna)
                device.write('SENS%s:FREQ:START %f' % (3,1250e6))
                device.write('SENS%s:FREQ:STOP %f' % (3,750e6))
                device.close()
                rm.close()

                for step2 in range(5):
                    if self.stop_event.is_set():
                        return
                    win_num = 3*step2+2
                    self.phase_def(win_num,2)

                for step3 in range(5):
                    if self.stop_event.is_set():
                        return
                    win_num = 3*step3+3
                    self.phase_def(win_num,4)

                for step4 in range(5):
                    if self.stop_event.is_set():
                        return
                    
                    pop = 'OK'
                    
                    if pop == 'OK':
                        win_num2 = 3*step4+1
                        win_num1 = 3*step4+2
                        self.phase_ref_qsrx(current_band[step4+i*5][2],2,win_num1)
                        win_num = 3*step4+1
                        time.sleep(0.5)
                        self.gain_ref_qsrx(current_band[step4+i*5][2],win_num2)

                    
                
                rm = visa.ResourceManager()
                device = rm.open_resource(ip_zna)
                device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx\\%s\\%s-%sGHz.znx'" % (dt_string,bands[i][0],bands[i][1]))
                time.sleep(2)
                #device.write(':SYSTem:PRESet')
                device.close()
                rm.close()
            
            rm = visa.ResourceManager()
            device = rm.open_resource(ip_zna)
            device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx\\%s\\6-8.5GHz.znx'" % (dt_string))
            device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx\\%s\\8.5-11GHz.znx'" % (dt_string))
            device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx\\%s\\11-13.5GHz.znx'" % (dt_string))
            device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx\\%s\\13.5-15.5GHz.znx'" % (dt_string))
            device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx\\%s\\15.5-18GHz.znx'" % (dt_string))
            time.sleep(2)
            #device.write(':SYSTem:PRESet')
            device.close()
            rm.close()
            

            for chan in [3,4]:
                if self.stop_event.is_set():
                    return

                chan_change = self.mbox("Connect VNA Port-1 to CH-%s input of Band 3\nConnect VNA Port-2 to CH-%s DRx Output" % (chan,chan))
                if chan_change == 'OK':
                    #Window 1
                    for j in range(5):
                        if self.stop_event.is_set():
                            return
                        rm = visa.ResourceManager()
                        device = rm.open_resource(ip_zna)
                        device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx\\%s\\%s-%sGHz.znx'" % (dt_string,bands[j][0],bands[j][1]))
                        time.sleep(5)
                        device.close()
                        rm.close()
                        for step6 in range(5):
                            if self.stop_event.is_set():
                                return
                            
                            pop = 'OK'

                            if pop == 'OK':
                                
                                win1 = 3*step6+1
                                self.new_gain_qsrx(current_band[step6+j*5][2],win1,chan)
                                time.sleep(1)
                                win2 = 3*step6+2
                                self.new_phase_qsrx(current_band[step6+j*5][2],win2,2,chan)
                        
                        rm = visa.ResourceManager()
                        device = rm.open_resource(ip_zna)
                        device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx\\%s\\%s-%sGHz.znx'" % (dt_string,bands[j][0],bands[j][1]))
                        time.sleep(5)
                        device.close()

                else:
                    self.stop_task()
                    break
                

            ref_change = self.mbox("Connect VNA Port-1 to CH-1 Input of Band 3\nConnect VNA Port-2 to CH-1 DRx Output\nConnect VNA Ref Port to CH-4 Input of Band 3\nConnect VNA Port-4 to CH-4 DRx Output")

            if ref_change == 'OK':
                
                for k in range(5):
                    if self.stop_event.is_set():
                        return

                    rm = visa.ResourceManager()
                    device = rm.open_resource(ip_zna)
                    device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx\\%s\\%s-%sGHz.znx'" % (dt_string,bands[k][0],bands[k][1]))
                    time.sleep(5)
                    device.close()
                    rm.close()

                    for step7 in range(5):
                        if self.stop_event.is_set():
                            return
                        
                        pop = 'OK'

                        if pop == 'OK':
                            win_num = 3*step7+3
                            self.phase_ref_qsrx(current_band[step7+k*5][2],4,win_num)
                    
                    rm = visa.ResourceManager()
                    device = rm.open_resource(ip_zna)
                    device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx\\%s\\%s-%sGHz.znx'" % (dt_string,bands[k][0],bands[k][1]))
                    time.sleep(5)
                    device.close()
                    rm.close()



                for chan in [2,3]:
                    if self.stop_event.is_set():
                        return
                    chan_change = self.mbox("Connect VNA Port-1 to CH-%s Input of Band 3\nConnect VNA Port-2 to CH-%s DRx Output" % (chan,chan))
                    if chan_change == 'OK':
                        #Window 1
                        for l in range(5):
                            if self.stop_event.is_set():
                                return
                            rm = visa.ResourceManager()
                            device = rm.open_resource(ip_zna)
                            device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx\\%s\\%s-%sGHz.znx'" % (dt_string,bands[l][0],bands[l][1]))
                            time.sleep(5)
                            device.close()
                            rm.close()
                            for step8 in range(5):
                                if self.stop_event.is_set():
                                    return
                                
                                pop = 'OK'

                                if pop == 'OK':
                                    win1 = 3*step8+1
                                    if chan == 2:
                                        self.new_gain_qsrx(current_band[step8+l*5][2],win1,chan)
                                    win2 = 3*step8+3
                                    self.new_phase_qsrx(current_band[step8+l*5][2],win2,4,chan)

                            rm = visa.ResourceManager()
                            device = rm.open_resource(ip_zna)
                            device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx\\%s\\%s-%sGHz.znx'" % (dt_string,bands[l][0],bands[l][1]))
                            time.sleep(5)
                            device.close()

                    else:
                        self.stop_task()
                        break
            else:
                self.stop_task()
            
            x,y,z = [],[],[]

            for m in range(5):
                rm = visa.ResourceManager()
                device = rm.open_resource(ip_zna)
                device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx\\%s\\%s-%sGHz.znx'" % (dt_string,bands[m][0],bands[m][1]))
                time.sleep(2)
                #device.write(':SYSTem:PRESet')
                device.close()
                rm.close()

                x1,y1,z1 = self.marker_qsr(5)

                x.extend(x1)
                y.extend(y1)
                z.extend(z1)

                rm = visa.ResourceManager()
                device = rm.open_resource(ip_zna)
                device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx\\%s\\%s-%sGHz.znx'" % (dt_string,bands[m][0],bands[m][1]))
                time.sleep(2)
                #device.write(':SYSTem:PRESet')
                device.close()
                rm.close()

            return x,y,z
        else:
            self.stop_task()

    def qsrx_band1_int(self,dt_string,bite='RF'):
        ip_zna = self.ip_zna

        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()

        qsrx = [[0.5,0.8,14.75,15],
                [0.7,1.2,14.55,14.6],
                [1.1,1.6,14.15,13.8],
                [1.5,2,13.75,13]]
        
        controls = ["1 1","1 0","0 1","0 0"]

        current_band = qsrx

        #ref_1 = self.mbox("Set Ref Port to %s \nSet to Channel 1\nCurrent Band is %s GHz - %s GHz" % (2,0.5,2.2))
        if bite=='BITE':
            _ = self.mbox("Set BITE control to 1")

        ref_1 = self.mbox("Current Band is %s GHz - %s GHz\nConnect VNA Port-1 to CH-1 Input of SFB\nConnect VNA Port-2 to CH-1 DRx Output\nConnect VNA Ref Port to CH-2 Input of SFB\nConnect VNA Port-4 to CH-2 DRx Output\nSet QSRx Band-1 Control : 1 1" % (0.5,2.2))
        
        if ref_1 == 'OK':
            #Window 1
            for step1 in range(4):
                if self.stop_event.is_set():
                    return
                
                self.mixer_settings_gain(current_band[step1],3*step1+1)
                self.mixer_settings_phase(2,current_band[step1],3*step1+2)
                self.mixer_settings_phase(4,current_band[step1],3*step1+3)
                
            time.sleep(2)

            rm = visa.ResourceManager()
            device = rm.open_resource(ip_zna)
            device.write('SENS%s:FREQ:START %f' % (2,750e6))
            device.write('SENS%s:FREQ:STOP %f' % (2,1050e6))
            device.write('SENS%s:FREQ:START %f' % (3,750e6))
            device.write('SENS%s:FREQ:STOP %f' % (3,1050e6))
            device.close()
            rm.close()

            for step2 in range(4):
                if self.stop_event.is_set():
                    return
                win_num = 3*step2+2
                self.phase_def(win_num,2)

            for step3 in range(4):
                if self.stop_event.is_set():
                    return
                win_num = 3*step3+3
                self.phase_def(win_num,4)

            for step4 in range(4):
                if self.stop_event.is_set():
                    return
                pop = self.mbox("Set SFB1 Band %s Control: %s" % (step4+1,controls[step4]))
                if pop == 'OK':
                    win_num1 = 3*step4+2
                    win_num2 = 3*step4+1
                    self.phase_ref_qsrx(current_band[step4][2],2,win_num1)
                    self.gain_ref_qsrx(current_band[step4][2],win_num2)
                else:
                    self.stop_task()
            
            rm = visa.ResourceManager()
            device = rm.open_resource(ip_zna)
            device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\0.5-2.2GHz.znx'" % (bite,dt_string))
            time.sleep(2)
            #device.write(':SYSTem:PRESet')
            device.close()
            rm.close()

            for chan in [3,4]:
                if self.stop_event.is_set():
                    return

                chan_change = self.mbox("Connect VNA Port-1 to CH-%s input of SFB\nConnect VNA Port-2 to CH-%s DRx Output" % (chan,chan))
                if chan_change == 'OK':
                    #Window 1
                    
                    for step5 in range(4):
                        if self.stop_event.is_set():
                            return
                        pop = self.mbox("Set SFB1 Band %s Control: %s" % (step5+1,controls[step5]))
                        if pop == 'OK':
                            win1 = 3*step5+1
                            self.new_gain_qsrx(current_band[step5][2],win1,chan)
                            win2 = 3*step5+2
                            self.new_phase_qsrx(current_band[step5][2],win2,2,chan)
                        else:
                            self.stop_task()
                else:
                    break

            ref_change = self.mbox("Connect VNA Port-1 to CH-1 Input of SFB\nConnect VNA Port-2 to CH-1 DRx Output\nConnect VNA Ref Port to CH-4 Input of SFB\nConnect VNA Port-4 to CH-4 DRx Output")

            if ref_change == 'OK':

                for step6 in range(4):
                    if self.stop_event.is_set():
                        return
                    pop = self.mbox("Set SFB1 Band %s Control: %s" % (step6+1,controls[step6]))
                    if pop == 'OK':
                        win_num = 3*step6+3
                        self.phase_ref_qsrx(current_band[step6][2],4,win_num)
                    else:
                        self.stop_task()


                for chan in [2,3]:
                    if self.stop_event.is_set():
                        return
                    chan_change = self.mbox("Connect VNA Port-1 to CH-%s input of SFB\nConnect VNA Port-2 to CH-%s DRx Output" % (chan,chan))
                    if chan_change == 'OK':

                        for step7 in range(4):
                            if self.stop_event.is_set():
                                return
                            pop = self.mbox("Set SFB1 Band %s Control: %s" % (step7+1,controls[step7]))
                            if pop == 'OK':
                                win1 = 3*step7+1
                                if chan == 2:
                                    self.new_gain_qsrx(current_band[step7][2],win1,chan)
                                win2 = 3*step7+3
                                self.new_phase_qsrx(current_band[step7][2],win2,4,chan)
                            else:
                                self.stop_task()

                    else:
                        self.stop_task()
                        break

                x,y,z = self.marker_qsr(4)

                rm = visa.ResourceManager()
                device = rm.open_resource(ip_zna)
                device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\0.5-2.2GHz.znx'" % (bite,dt_string))
                time.sleep(2)
                device.close()
                rm.close()

                return x,y,z
            
            else:
                self.stop_task()

    def qsrx_band2_int(self,dt_string,bite='RF'):
        ip_zna = self.ip_zna

        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()

        qsrx = [[2,2.5,13.25,12],
                [2.5,3,12.75,11],
                [3,3.5,12.25,10],
                [3.5,4,11.75,9],
                [4,4.5,11.25,8],
                [4.5,5,10.75,7],
                [4.5,5,16.75,11],
                [5,5.5,10.25,6],
                [5.5,6,9.75,5]]
        
        controls = ["0 1","1 0","0 0","1 1"]

        current_band = qsrx

        if bite=='BITE':
            _ = self.mbox("Set BITE control to 1")

        ref_1 = self.mbox("Current Band is %s GHz - %s GHz\nConnect VNA Port-1 to CH-1 Input of SFB\nConnect VNA Port-2 to CH-1 DRx Output\nConnect VNA Ref Port to CH-2 Input of SFB\nConnect VNA Port-4 to CH-2 DRx Output\nSet QSRx Band-2 Control : 0 1" % (2,6.25))
        
        if ref_1 == 'OK':
            #Window 1
            for step1 in range(4):
                if self.stop_event.is_set():
                    return
                self.mixer_settings_gain(current_band[step1],3*step1+1)
                self.mixer_settings_phase(2,current_band[step1],3*step1+2)
                self.mixer_settings_phase(4,current_band[step1],3*step1+3)

            time.sleep(2)

            rm = visa.ResourceManager()
            device = rm.open_resource(ip_zna)
            device.write('SENS%s:FREQ:START %f' % (3,750e6))
            device.write('SENS%s:FREQ:STOP %f' % (3,1250e6))
            device.close()
            rm.close()

            for step2 in range(4):
                if self.stop_event.is_set():
                    return
                win_num = 3*step2+2
                self.phase_def(win_num,2)

            for step3 in range(4):
                if self.stop_event.is_set():
                    return
                win_num = 3*step3+3
                self.phase_def(win_num,4)

            for step4 in range(4):
                if self.stop_event.is_set():
                    return
                win_num1 = 3*step4+2
                win_num2 = 3*step4+1
                if step4 == 0:
                    pop4 = self.mbox("Set SFB2 Band 1 Control: %s" % (controls[0]))
                elif step4 == 1:
                    pop4 = self.mbox("Set SFB2 Band 2 Control: %s" % (controls[1]))
                elif step4 == 3:
                    pop4 = self.mbox("Set SFB2 Band 3 Control: %s" % (controls[2]))
                else:
                    pop4 = 'OK'
                
                if pop4 == 'OK':
                    self.phase_ref_qsrx(current_band[step4][2],2,win_num1)
                    time.sleep(2)
                    self.gain_ref_qsrx(current_band[step4][2],win_num2)
                else:
                    self.stop_event()

            rm = visa.ResourceManager()
            device = rm.open_resource(ip_zna)
            device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\2-4GHz.znx'" % (bite,dt_string))
            time.sleep(2)
            #device.write(':SYSTem:PRESet')
            device.close()
            rm.close()
            
            #Window 2
            rm = visa.ResourceManager()
            device = rm.open_resource(ip_zna)
            device.write(':SYSTem:PRESet')
            device.close()
            rm.close()

            for step1 in range(5):
                if self.stop_event.is_set():
                    return
                self.mixer_settings_gain(current_band[step1+4],3*step1+1)
                self.mixer_settings_phase(2,current_band[step1+4],3*step1+2)
                self.mixer_settings_phase(4,current_band[step1+4],3*step1+3)
                
            time.sleep(5)

            rm = visa.ResourceManager()
            device = rm.open_resource(ip_zna)
            device.write('SENS%s:FREQ:START %f' % (3,750e6))
            device.write('SENS%s:FREQ:STOP %f' % (3,1250e6))
            device.close()
            rm.close()

            for step2 in range(5):
                if self.stop_event.is_set():
                    return
                win_num = 3*step2+2
                self.phase_def(win_num,2)

            for step3 in range(5):
                if self.stop_event.is_set():
                    return
                win_num = 3*step3+3
                self.phase_def(win_num,4)

            for step4 in range(5):
                if self.stop_event.is_set():
                    return
                win_num1 = 3*step4+2
                win_num2 = 3*step4+1
                if step4 == 3:
                    pop4 = self.mbox("Set SFB2 Band 4 Control: %s" % (controls[3]))
                else:
                    pop4 = 'OK'

                if pop4 == 'OK':
                    self.phase_ref_qsrx(current_band[step4+4][2],2,win_num1)
                    time.sleep(2)
                    self.gain_ref_qsrx(current_band[step4+4][2],win_num2)
                else:
                    self.stop_task()

            rm = visa.ResourceManager()
            device = rm.open_resource(ip_zna)
            device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\4-6GHz.znx'" % (bite,dt_string))
            time.sleep(1)
            device.close()
            rm.close()

            rm = visa.ResourceManager()
            device = rm.open_resource(ip_zna)
            #device.write(':SYSTem:PRESet')
            device.close()
            rm.close()

            

            for chan in [3,4]:
                if self.stop_event.is_set():
                    return

                chan_change = self.mbox("Connect VNA Port-1 to CH-%s input of SFB\nConnect VNA Port-2 to CH-%s DRx Output" % (chan,chan))
                if chan_change == 'OK':
                    #Window 1
                    rm = visa.ResourceManager()
                    device = rm.open_resource(ip_zna)
                    device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\2-4GHz.znx'" % (bite,dt_string))
                    time.sleep(5)
                    device.close()
                    rm.close()
                    for step6 in range(4):
                        if self.stop_event.is_set():
                            return
                        if step6 == 0:
                            pop6 = self.mbox("Set SFB2 Band 1 Control: %s" % (controls[0]))
                        elif step6 == 1:
                            pop6 = self.mbox("Set SFB2 Band 2 Control: %s" % (controls[1]))
                        elif step6 == 3:
                            pop6 = self.mbox("Set SFB2 Band 3 Control: %s" % (controls[2]))
                        else:
                            pop6 = 'OK'
                        if pop6 == 'OK':
                            win1 = 3*step6+1
                            self.new_gain_qsrx(current_band[step6][2],win1,chan)
                            win2 = 3*step6+2
                            self.new_phase_qsrx(current_band[step6][2],win2,2,chan)
                        else:
                            self.stop_task()


                    #Window 2
                    rm = visa.ResourceManager()
                    device = rm.open_resource(ip_zna)
                    device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\2-4GHz.znx'" % (bite,dt_string))
                    time.sleep(2)
                    device.close()
                    rm.close()

                    rm = visa.ResourceManager()
                    device = rm.open_resource(ip_zna)
                    #device.write(':SYSTem:PRESet')
                    time.sleep(2)
                    device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\4-6GHz.znx'" % (bite,dt_string))
                    time.sleep(5)
                    device.close()
                    rm.close()
                    for step6 in range(5):
                        if self.stop_event.is_set():
                            return
                        win1 = 3*step6+1
                        if step6 == 3:
                            pop6 = self.mbox("Set SFB2 Band 4 Control: %s" % (controls[3]))
                        else:
                            pop6 = 'OK'

                        if pop6 == 'OK':
                            self.new_gain_qsrx(current_band[step6+4][2],win1,chan)
                            win2 = 3*step6+2
                            self.new_phase_qsrx(current_band[step6+4][2],win2,2,chan)
                        else:
                            self.stop_task()

                    rm = visa.ResourceManager()
                    device = rm.open_resource(ip_zna)
                    device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\4-6GHz.znx'" % (bite,dt_string))
                    time.sleep(2)
                    #device.write(':SYSTem:PRESet')
                    device.close()
                    rm.close()

                    rm = visa.ResourceManager()
                    device = rm.open_resource(ip_zna)
                    #device.write(':SYSTem:PRESet')
                    device.close()
                    rm.close()
                
                else:
                    self.stop_task()
                    break
                



            ref_change = self.mbox("Connect VNA Port-1 to CH-1 Input of SFB\nConnect VNA Port-2 to CH-1 DRx Output\nConnect VNA Ref Port to CH-4 Input of SFB\nConnect VNA Port-4 to CH-4 DRx Output")

            if ref_change == 'OK':
                rm = visa.ResourceManager()
                device = rm.open_resource(ip_zna)
                device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\2-4GHz.znx'" % (bite,dt_string))
                time.sleep(5)
                device.close()
                rm.close()

                for step7 in range(4):
                    if self.stop_event.is_set():
                        return
                    win_num = 3*step7+3
                    if step7 == 0:
                        pop7 = self.mbox("Set SFB2 Band 1 Control: %s" % (controls[0]))
                    elif step7 == 1:
                        pop7 = self.mbox("Set SFB2 Band 2 Control: %s" % (controls[1]))
                    elif step7 == 3:
                        pop7 = self.mbox("Set SFB2 Band 3 Control: %s" % (controls[2]))
                    else:
                        pop7 = 'OK'

                    if pop7 == 'OK':
                        self.phase_ref_qsrx(current_band[step7][2],4,win_num)
                    else:
                        self.stop_task()

                rm = visa.ResourceManager()
                device = rm.open_resource(ip_zna)
                device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\2-4GHz.znx'" % (bite,dt_string))
                time.sleep(2)
                device.close()
                rm.close()

                rm = visa.ResourceManager()
                device = rm.open_resource(ip_zna)
                #device.write(':SYSTem:PRESet')
                time.sleep(2)
                device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\4-6GHz.znx'" % (bite,dt_string))
                time.sleep(5)
                device.close()
                rm.close()

                for step7 in range(5):
                    if self.stop_event.is_set():
                        return
                    if step7 == 3 :
                        pop7 = self.mbox("Set SFB2 Band 4 Control: %s" % (controls[3]))
                    else:
                        pop7 = 'OK'
                    if pop7 == 'OK':
                        win_num = 3*step7+3
                        self.phase_ref_qsrx(current_band[step7+4][2],4,win_num)
                    else:
                        self.stop_task()

                rm = visa.ResourceManager()
                device = rm.open_resource(ip_zna)
                device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\4-6GHz.znx'" % (bite,dt_string))
                time.sleep(2)
                device.close()
                rm.close()

                for chan in [2,3]:
                    if self.stop_event.is_set():
                        return
                    chan_change = self.mbox("Connect VNA Port-1 to CH-%s input of SFB\nConnect VNA Port-2 to CH-%s DRx Output" % (chan,chan))
                    if chan_change == 'OK':
                        #Window 1
                        rm = visa.ResourceManager()
                        device = rm.open_resource(ip_zna)
                        device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\2-4GHz.znx'" % (bite,dt_string))
                        time.sleep(5)
                        device.close()
                        rm.close()
                        for step8 in range(4):
                            win1 = 3*step8+1
                            if step8 == 0:
                                pop8 = self.mbox("Set SFB2 Band 1 Control: %s" % (controls[0]))
                            elif step8 == 1:
                                pop8 = self.mbox("Set SFB2 Band 2 Control: %s" % (controls[1]))
                            elif step8 == 3:
                                pop8 = self.mbox("Set SFB2 Band 3 Control: %s" % (controls[2]))
                            else:
                                pop8 = 'OK'
                                
                            if pop8 == 'OK':
                                if chan == 2:
                                    self.new_gain_qsrx(current_band[step8][2],win1,chan)
                                win2 = 3*step8+3
                                self.new_phase_qsrx(current_band[step8][2],win2,4,chan)
                            else:
                                self.stop_task()


                        #Window 2
                        rm = visa.ResourceManager()
                        device = rm.open_resource(ip_zna)
                        device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\2-4GHz.znx'" % (bite,dt_string))
                        time.sleep(2)
                        device.close()
                        rm.close()

                        rm = visa.ResourceManager()
                        device = rm.open_resource(ip_zna)
                        #device.write(':SYSTem:PRESet')
                        time.sleep(2)
                        device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\4-6GHz.znx'" % (bite,dt_string))
                        time.sleep(5)
                        device.close()
                        rm.close()
                        for step8 in range(5):
                            if self.stop_event.is_set():
                                return
                            win1 = 3*step8+1
                            if step8 == 3:
                                pop8 = self.mbox("Set SFB2 Band 4 Control: %s" % (controls[3]))
                            else:
                                pop8 = 'OK'

                            if pop8 == 'OK':
                                if chan == 2:
                                    self.new_gain_qsrx(current_band[step8+4][2],win1,chan)
                                win2 = 3*step8+3
                                self.new_phase_qsrx(current_band[step8+4][2],win2,4,chan)
                            else:
                                self.stop_task()

                        rm = visa.ResourceManager()
                        device = rm.open_resource(ip_zna)
                        device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\4-6GHz.znx'" % (bite,dt_string))
                        time.sleep(2)
                        #device.write(':SYSTem:PRESet')
                        device.close()
                        rm.close()
                    else:
                        self.stop_task()
                        break
            
            rm = visa.ResourceManager()
            device = rm.open_resource(ip_zna)
            device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\2-4GHz.znx'" % (bite,dt_string))
            time.sleep(2)
            #device.write(':SYSTem:PRESet')
            device.close()
            rm.close()

            x1,y1,z1 = self.marker_qsr(4)

            rm = visa.ResourceManager()
            device = rm.open_resource(ip_zna)
            device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\2-4GHz.znx'" % (bite,dt_string))
            time.sleep(2)
            #device.write(':SYSTem:PRESet')
            device.close()
            rm.close()

            rm = visa.ResourceManager()
            device = rm.open_resource(ip_zna)
            device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\4-6GHz.znx'" % (bite,dt_string))
            time.sleep(2)
            #device.write(':SYSTem:PRESet')
            device.close()
            rm.close()

            x2,y2,z2 = self.marker_qsr(5)

            rm = visa.ResourceManager()
            device = rm.open_resource(ip_zna)
            device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\4-6GHz.znx'" % (bite,dt_string))
            time.sleep(2)
            #device.write(':SYSTem:PRESet')
            device.close()
            rm.close()

            x1.extend(x2)
            y1.extend(y2)
            z1.extend(z2)

            return x1,y1,z1
        
        else:
            self.stop_task()

    def qsrx_band3_int(self,dt_string,bite='RF'):
        ip_zna = self.ip_zna

        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()

        qsrx = [[6,6.5,9.25,4],
                [6.5,7,9.75,4],
                [7,7.5,10.25,4],
                [7.5,8,10.75,4],
                [8,8.5,11.25,4],
                [8.5,9,11.75,4],
                [9,9.5,12.25,4],
                [9.5,10,12.75,4],
                [10,10.5,13.25,4],
                [10.5,11,13.75,4],
                [11,11.5,14.25,4],
                [11.5,12,14.75,4],
                [12,12.5,9.25,4],
                [12.5,13,9.75,4],
                [13,13.5,10.25,4],
                [13.5,14,10.75,4],
                [13.5,14,16.75,4],
                [14,14.5,11.25,4],
                [14.5,15,11.75,4],
                [15,15.5,12.25,4],
                [15.5,16,12.75,4],
                [16,16.5,13.25,4],
                [16.5,17,13.75,4],
                [17,17.5,14.25,4],
                [17.5,18,14.75,4]]
        
        controls = ["0 0","0 1","1 0","1 1"]

        current_band = qsrx

        bands = [[6,8.5],
                [8.5,11],
                [11,13.5],
                [13.5,15.5],
                [15.5,18]]

        if bite == 'BITE':
            _ = self.mbox("Set BITE control to 1")

        ref_1 = self.mbox("Current Band is %s GHz - %s GHz\nConnect VNA Port-1 to CH-1 Input of SFB\nConnect VNA Port-2 to CH-1 DRx Output\nConnect VNA Ref Port to CH-2 Input of SFB\nConnect VNA Port-4 to CH-2 DRx Output\nSet QSRx Band-2 Control : 0 0" % (6,18))
        

        if ref_1 == 'OK':
            #Window 1
            for i in range(5):
                if self.stop_event.is_set():
                    return
                rm = visa.ResourceManager()
                device = rm.open_resource(ip_zna)
                device.write(':SYSTem:PRESet')
                device.close()
                rm.close()

                for step1 in range(5):
                    if self.stop_event.is_set():
                        return
                    self.mixer_settings_gain(current_band[step1+i*5],3*step1+1)
                    self.mixer_settings_phase(2,current_band[step1+i*5],3*step1+2)
                    self.mixer_settings_phase(4,current_band[step1+i*5],3*step1+3)
                    
                time.sleep(2)

                rm = visa.ResourceManager()
                device = rm.open_resource(ip_zna)
                device.write('SENS%s:FREQ:START %f' % (3,1250e6))
                device.write('SENS%s:FREQ:STOP %f' % (3,750e6))
                device.close()
                rm.close()

                for step2 in range(5):
                    if self.stop_event.is_set():
                        return
                    win_num = 3*step2+2
                    self.phase_def(win_num,2)

                for step3 in range(5):
                    if self.stop_event.is_set():
                        return
                    win_num = 3*step3+3
                    self.phase_def(win_num,4)

                for step4 in range(5):
                    if self.stop_event.is_set():
                        return
                    if i == 0 and step4 == 0:
                        pop = self.mbox("Set SFB3 Band 1 Control: %s" % (controls[0]))
                    elif i == 1 and step4 == 2:
                        pop = self.mbox("Set SFB3 Band 2 Control: %s" % (controls[1]))
                    elif i == 2 and step4 == 4:
                        pop = self.mbox("Set SFB3 Band 3 Control: %s" % (controls[2]))
                    elif i == 3 and step4 == 3:
                        pop = self.mbox("Set SFB3 Band 4 Control: %s" % (controls[3]))
                    else:
                        pop = 'OK'
                    
                    if pop == 'OK':
                        win_num2 = 3*step4+1
                        win_num1 = 3*step4+2
                        self.phase_ref_qsrx(current_band[step4+i*5][2],2,win_num1)
                        win_num = 3*step4+1
                        time.sleep(0.5)
                        self.gain_ref_qsrx(current_band[step4+i*5][2],win_num2)
                    else:
                        self.stop_task()

                    
                
                rm = visa.ResourceManager()
                device = rm.open_resource(ip_zna)
                device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\%s-%sGHz.znx'" % (bite,dt_string,bands[i][0],bands[i][1]))
                time.sleep(2)
                #device.write(':SYSTem:PRESet')
                device.close()
                rm.close()
            
            rm = visa.ResourceManager()
            device = rm.open_resource(ip_zna)
            device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\6-8.5GHz.znx'" % (bite,dt_string))
            device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\8.5-11GHz.znx'" % (bite,dt_string))
            device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\11-13.5GHz.znx'" % (bite,dt_string))
            device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\13.5-15.5GHz.znx'" % (bite,dt_string))
            device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\15.5-18GHz.znx'" % (bite,dt_string))
            time.sleep(2)
            #device.write(':SYSTem:PRESet')
            device.close()
            rm.close()
            

            for chan in [3,4]:
                if self.stop_event.is_set():
                    return

                chan_change = self.mbox("Connect VNA Port-1 to CH-%s input of SFB\nConnect VNA Port-2 to CH-%s DRx Output" % (chan,chan))
                if chan_change == 'OK':
                    #Window 1
                    for j in range(5):
                        if self.stop_event.is_set():
                            return
                        rm = visa.ResourceManager()
                        device = rm.open_resource(ip_zna)
                        device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\%s-%sGHz.znx'" % (bite,dt_string,bands[j][0],bands[j][1]))
                        time.sleep(5)
                        device.close()
                        rm.close()
                        for step6 in range(5):
                            if self.stop_event.is_set():
                                return
                            if j == 0 and step6 == 0:
                                pop = self.mbox("Set SFB3 Band 1 Control: %s" % (controls[0]))
                            elif j == 1 and step6 == 2:
                                pop = self.mbox("Set SFB3 Band 2 Control: %s" % (controls[1]))
                            elif j == 2 and step6 == 4:
                                pop = self.mbox("Set SFB3 Band 3 Control: %s" % (controls[2]))
                            elif j == 3 and step6 == 3:
                                pop = self.mbox("Set SFB3 Band 4 Control: %s" % (controls[3]))
                            else:
                                pop = 'OK'

                            if pop == 'OK':
                                
                                win1 = 3*step6+1
                                self.new_gain_qsrx(current_band[step6+j*5][2],win1,chan)
                                time.sleep(1)
                                win2 = 3*step6+2
                                self.new_phase_qsrx(current_band[step6+j*5][2],win2,2,chan)
                            else:
                                self.stop_task()
                        
                        rm = visa.ResourceManager()
                        device = rm.open_resource(ip_zna)
                        device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\%s-%sGHz.znx'" % (bite,dt_string,bands[j][0],bands[j][1]))
                        time.sleep(5)
                        device.close()

                else:
                    self.stop_task()
                    break
                

            ref_change = self.mbox("Connect VNA Port-1 to CH-1 Input of SFB\nConnect VNA Port-2 to CH-1 DRx Output\nConnect VNA Ref Port to CH-4 Input of SFB\nConnect VNA Port-4 to CH-4 DRx Output")

            if ref_change == 'OK':
                
                for k in range(5):
                    if self.stop_event.is_set():
                        return

                    rm = visa.ResourceManager()
                    device = rm.open_resource(ip_zna)
                    device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\%s-%sGHz.znx'" % (bite,dt_string,bands[k][0],bands[k][1]))
                    time.sleep(5)
                    device.close()
                    rm.close()

                    for step7 in range(5):
                        if self.stop_event.is_set():
                            return
                        if k == 0 and step7 == 0:
                            pop = self.mbox("Set SFB3 Band 1 Control: %s" % (controls[0]))
                        elif k == 1 and step7 == 2:
                            pop = self.mbox("Set SFB3 Band 2 Control: %s" % (controls[1]))
                        elif k == 2 and step7 == 4:
                            pop = self.mbox("Set SFB3 Band 3 Control: %s" % (controls[2]))
                        elif k == 3 and step7 == 3:
                            pop = self.mbox("Set SFB3 Band 4 Control: %s" % (controls[3]))
                        else:
                            pop = 'OK'

                        if pop == 'OK':
                            win_num = 3*step7+3
                            self.phase_ref_qsrx(current_band[step7+k*5][2],4,win_num)
                        else:
                            self.stop_task()
                    
                    rm = visa.ResourceManager()
                    device = rm.open_resource(ip_zna)
                    device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\%s-%sGHz.znx'" % (bite,dt_string,bands[k][0],bands[k][1]))
                    time.sleep(5)
                    device.close()
                    rm.close()



                for chan in [2,3]:
                    if self.stop_event.is_set():
                        return
                    chan_change = self.mbox("Connect VNA Port-1 to CH-%s input of SFB\nConnect VNA Port-2 to CH-%s DRx Output" % (chan,chan))
                    if chan_change == 'OK':
                        #Window 1
                        for l in range(5):
                            if self.stop_event.is_set():
                                return
                            rm = visa.ResourceManager()
                            device = rm.open_resource(ip_zna)
                            device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\%s-%sGHz.znx'" % (bite,dt_string,bands[l][0],bands[l][1]))
                            time.sleep(5)
                            device.close()
                            rm.close()
                            for step8 in range(5):
                                if self.stop_event.is_set():
                                    return
                                if l == 0 and step8 == 0:
                                    pop = self.mbox("Set SFB3 Band 1 Control: %s" % (controls[0]))
                                elif l == 1 and step8 == 2:
                                    pop = self.mbox("Set SFB3 Band 2 Control: %s" % (controls[1]))
                                elif l == 2 and step8 == 4:
                                    pop = self.mbox("Set SFB3 Band 3 Control: %s" % (controls[2]))
                                elif l == 3 and step8 == 3:
                                    pop = self.mbox("Set SFB3 Band 4 Control: %s" % (controls[3]))
                                else:
                                    pop = 'OK'

                                if pop == 'OK':
                                    win1 = 3*step8+1
                                    if chan == 2:
                                        self.new_gain_qsrx(current_band[step8+l*5][2],win1,chan)
                                    win2 = 3*step8+3
                                    self.new_phase_qsrx(current_band[step8+l*5][2],win2,4,chan)
                                else:
                                    self.stop_task()

                            rm = visa.ResourceManager()
                            device = rm.open_resource(ip_zna)
                            device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\%s-%sGHz.znx'" % (bite,dt_string,bands[l][0],bands[l][1]))
                            time.sleep(5)
                            device.close()

                    else:
                        self.stop_task()
                        break
            
                x,y,z = [],[],[]

                for m in range(5):
                    rm = visa.ResourceManager()
                    device = rm.open_resource(ip_zna)
                    device.write("MMEM:LOAD:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\%s-%sGHz.znx'" % (bite,dt_string,bands[m][0],bands[m][1]))
                    time.sleep(2)
                    #device.write(':SYSTem:PRESet')
                    device.close()
                    rm.close()

                    x1,y1,z1 = self.marker_qsr(5)

                    x.extend(x1)
                    y.extend(y1)
                    z.extend(z1)

                    rm = visa.ResourceManager()
                    device = rm.open_resource(ip_zna)
                    device.write("MMEM:STOR:STAT 1,'C:\\Synergy\\QSRx_int\\%s\\%s\\%s-%sGHz.znx'" % (bite,dt_string,bands[m][0],bands[m][1]))
                    time.sleep(2)
                    #device.write(':SYSTem:PRESet')
                    device.close()
                    rm.close()

                return x,y,z
            else:
                self.stop_task()
        
        else:
            self.stop_task()

    def data_trace_off_qsr(self):
        ip_zna = self.ip_zna

        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)

        chans = str(device.query("CONF:CHAN:CAT?"))
        l1 = chans.split(",")
        iter = float(l1[-2])/3

        for i in range(int(iter)):
            device.write("DISP:TRAC:SHOW 'Gain%s', OFF" % (3*i+1))
            device.write("DISP:TRAC:SHOW 'Gain%s_mem', OFF" % (3*i+1))
            time.sleep(0.5)
            device.write("DISP:TRAC:SHOW 'Phase%s_R2', OFF" % (3*i+2))
            device.write("DISP:TRAC:SHOW 'Phase%s_R2_mem', OFF" % (3*i+2))
            time.sleep(0.5)
            device.write("DISP:TRAC:SHOW 'Phase%s_R4', OFF" % (3*i+3))
            device.write("DISP:TRAC:SHOW 'Phase%s_R4_mem', OFF" % (3*i+3))
            time.sleep(0.5)


        device.close()
        rm.close()

    def track_qsr(self):
        df = self.make_df()
        
        dut = self.qsrx_sl

        rm = visa.ResourceManager()
        device = rm.open_resource(self.ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()

        folder = self.new_folder(dut,False,False)
        #qsrx_band3(folder)
        a1,p21,p41 = self.qsrx_band1(folder)

        for i in range(4):
            df.iloc[i,4] = p21[i]
            df.iloc[i,5] = p41[i]
            df.iloc[i,6] = max(p21[i],p41[i])
            df.iloc[i,7] = a1[i]

        a2,p22,p42 = self.qsrx_band2(folder)

        for j in range(9):
            df.iloc[4+j,4] = p22[j]
            df.iloc[4+j,5] = p42[j]
            df.iloc[4+j,6] = max(p22[j],p42[j])
            df.iloc[4+j,7] = a2[j]

        a3,p23,p43 = self.qsrx_band3(folder)

        self.ps_output(0)

        for k in range(25):
            df.iloc[13+k,4] = p23[k]
            df.iloc[13+k,5] = p43[k]
            df.iloc[13+k,6] = max(p23[k],p43[k])
            df.iloc[13+k,7] = a3[k]
        
        df.to_excel("test_data/%s/%s_QSRx.xlsx" % (folder,folder),index_label="Band")
        
        self.report_qsr(folder,"template_RF.docx","report_QSRx.docx","report_QSRx.pdf",df)

    def track_qsr_band1(self):
        df = self.make_df()
        
        dut = self.qsrx_sl

        rm = visa.ResourceManager()
        device = rm.open_resource(self.ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()

        folder = self.new_folder(dut,False,False)
        #qsrx_band3(folder)
        a1,p21,p41 = self.qsrx_band1(folder)
        
        self.ps_output(0)

        for i in range(4):
            df.iloc[i,4] = p21[i]
            df.iloc[i,5] = p41[i]
            df.iloc[i,6] = max(p21[i],p41[i])
            df.iloc[i,7] = a1[i]

        df.to_excel("test_data/%s/%s_QSRx_Band1.xlsx" % (folder,folder),index_label="Band")
        
        self.report_qsr(folder,"template_RF.docx","report_QSRx_Band1.docx","report_QSRx_Band1.pdf",df)

    def track_qsr_band2(self):
        df = self.make_df()
        
        dut = self.qsrx_sl

        rm = visa.ResourceManager()
        device = rm.open_resource(self.ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()

        folder = self.new_folder(dut,False,False)

        a2,p22,p42 = self.qsrx_band2(folder)

        self.ps_output(0)

        for j in range(9):
            df.iloc[4+j,4] = p22[j]
            df.iloc[4+j,5] = p42[j]
            df.iloc[4+j,6] = max(p22[j],p42[j])
            df.iloc[4+j,7] = a2[j]
        
        df.to_excel("test_data/%s/%s_QSRx_Band2.xlsx" % (folder,folder),index_label="Band")
        
        self.report_qsr(folder,"template_RF.docx","report_QSRx_Band2.docx","report_QSRx_Band2.pdf",df)

    def track_qsr_band3(self):
        df = self.make_df()
        
        dut = self.qsrx_sl

        rm = visa.ResourceManager()
        device = rm.open_resource(self.ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()

        folder = self.new_folder(dut,False,False)
        a3,p23,p43 = self.qsrx_band3(folder)

        self.ps_output(0)

        for k in range(25):
            df.iloc[13+k,4] = p23[k]
            df.iloc[13+k,5] = p43[k]
            df.iloc[13+k,6] = max(p23[k],p43[k])
            df.iloc[13+k,7] = a3[k]
        
        df.to_excel("test_data/%s/%s_QSRx_Band3.xlsx" % (folder,folder),index_label="Band")
        
        self.report_qsr(folder,"template_RF.docx","report_QSRx_Band3.docx","report_QSRx_Band3.pdf",df)

    def track_qsr_int_rf(self):
        df = self.make_df()
        
        dut = self.sfb1_sl+"_"+self.sfb2_sl+"_"+self.sfb3_sl+"_"+self.qsrx_sl

        rm = visa.ResourceManager()
        device = rm.open_resource(self.ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()

        folder = self.new_folder(dut,True,False)
        
        
        a1,p21,p41 = self.qsrx_band1_int(folder,'RF')

        for i in range(4):
            df.iloc[i,4] = p21[i]
            df.iloc[i,5] = p41[i]
            df.iloc[i,6] = max(p21[i],p41[i])
            df.iloc[i,7] = a1[i]

        self.ps_output(0)

        sf_change = self.mbox("Change module to SFB2")

        if sf_change == 'OK':
            self.ps_output(1)
            a2,p22,p42 = self.qsrx_band2_int(folder,'RF')

        for j in range(9):
            df.iloc[4+j,4] = p22[j]
            df.iloc[4+j,5] = p42[j]
            df.iloc[4+j,6] = max(p22[j],p42[j])
            df.iloc[4+j,7] = a2[j]
        
        self.ps_output(0)

        sf_change = self.mbox("Change module to SFB3")

        if sf_change == 'OK':
            self.ps_output(1)
            a3,p23,p43 = self.qsrx_band3_int(folder,'RF')

        self.ps_output(0)

        for k in range(25):
            df.iloc[13+k,4] = p23[k]
            df.iloc[13+k,5] = p43[k]
            df.iloc[13+k,6] = max(p23[k],p43[k])
            df.iloc[13+k,7] = a3[k]
        
        df.to_excel("test_data/%s/%s_QSRx_integrated_RF.xlsx" % (folder,folder),index_label="Band")
        
        self.report_qsr(folder,"template_RF.docx","report_QSRx_Integrated_RF.docx","report_QSRx_Integrated_RF.pdf",df)

    def track_qsr_sfb1_rf(self):
        df = self.make_df()
        
        dut = self.sfb1_sl+self.qsrx_sl

        rm = visa.ResourceManager()
        device = rm.open_resource(self.ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()

        folder = self.new_folder(dut,True,False)
        
        
        a1,p21,p41 = self.qsrx_band1_int(folder,'RF')

        for i in range(4):
            df.iloc[i,4] = p21[i]
            df.iloc[i,5] = p41[i]
            df.iloc[i,6] = max(p21[i],p41[i])
            df.iloc[i,7] = a1[i]

        self.ps_output(0)
        
        df.to_excel("test_data/%s/%s_QSRx_integrated_SFB1_RF.xlsx" % (folder,folder),index_label="Band")
        
        self.report_qsr(folder,"template_RF.docx","report_QSRx_Integrated_SFB1_RF.docx","report_QSRx_Integrated_SFB1_RF.pdf",df)

    def track_qsr_sfb2_rf(self):
        df = self.make_df()
        
        dut = self.sfb2_sl+self.qsrx_sl

        rm = visa.ResourceManager()
        device = rm.open_resource(self.ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()

        folder = self.new_folder(dut,True,False)
        
        
        a2,p22,p42 = self.qsrx_band2_int(folder,'RF')

        for j in range(9):
            df.iloc[4+j,4] = p22[j]
            df.iloc[4+j,5] = p42[j]
            df.iloc[4+j,6] = max(p22[j],p42[j])
            df.iloc[4+j,7] = a2[j]

        self.ps_output(0)
        
        df.to_excel("test_data/%s/%s_QSRx_integrated_SFB2_RF.xlsx" % (folder,folder),index_label="Band")
        
        self.report_qsr(folder,"template_RF.docx","report_QSRx_Integrated_SFB2_RF.docx","report_QSRx_Integrated_SFB2_RF.pdf",df)

    def track_qsr_sfb3_rf(self):
        df = self.make_df()
        
        dut = self.sfb3_sl+self.qsrx_sl

        rm = visa.ResourceManager()
        device = rm.open_resource(self.ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()

        folder = self.new_folder(dut,True,False)
        
        a3,p23,p43 = self.qsrx_band3_int(folder,'RF')

        for k in range(25):
            df.iloc[13+k,4] = p23[k]
            df.iloc[13+k,5] = p43[k]
            df.iloc[13+k,6] = max(p23[k],p43[k])
            df.iloc[13+k,7] = a3[k]

        self.ps_output(0)
        
        df.to_excel("test_data/%s/%s_QSRx_integrated_SFB3_RF.xlsx" % (folder,folder),index_label="Band")
        
        self.report_qsr(folder,"template_RF.docx","report_QSRx_Integrated_SFB3_RF.docx","report_QSRx_Integrated_SFB3_RF.pdf",df)

    def track_qsr_int_bite(self):
        df = self.make_df()
        
        dut = self.sfb1_sl+"_"+self.sfb2_sl+"_"+self.sfb3_sl+"_"+self.qsrx_sl

        rm = visa.ResourceManager()
        device = rm.open_resource(self.ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()

        folder = self.new_folder(dut,True,True)
        a1,p21,p41 = self.qsrx_band1_int(folder,'BITE')

        for i in range(4):
            df.iloc[i,4] = p21[i]
            df.iloc[i,5] = p41[i]
            df.iloc[i,6] = max(p21[i],p41[i])
            df.iloc[i,7] = a1[i]

        self.ps_output(0)

        sf_change = self.mbox("Change module to SFB2")

        if sf_change == 'OK':
            self.ps_output(1)
            a2,p22,p42 = self.qsrx_band2_int(folder,'BITE')

        for j in range(9):
            df.iloc[4+j,4] = p22[j]
            df.iloc[4+j,5] = p42[j]
            df.iloc[4+j,6] = max(p22[j],p42[j])
            df.iloc[4+j,7] = a2[j]

        self.ps_output(0)

        sf_change = self.mbox("Change module to SFB3")

        if sf_change == 'OK':
            self.ps_output(1)
            a3,p23,p43 = self.qsrx_band3_int(folder,'BITE')
        
        self.ps_output(0)

        for k in range(25):
            df.iloc[13+k,4] = p23[k]
            df.iloc[13+k,5] = p43[k]
            df.iloc[13+k,6] = max(p23[k],p43[k])
            df.iloc[13+k,7] = a3[k]
        
        df.to_excel("test_data/%s/%s_QSRx_integrated_BITE.xlsx" % (folder,folder),index_label="Band")
        
        self.report_qsr(folder,"template_BITE.docx","report_QSRx_Integrated_BITE.docx","report_QSRx_Integrated_BITE.pdf",df)

    def track_qsr_sfb1_bite(self):
        df = self.make_df()
        
        dut = self.sfb1_sl+"_"+self.qsrx_sl

        rm = visa.ResourceManager()
        device = rm.open_resource(self.ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()

        folder = self.new_folder(dut,True,True)
        a1,p21,p41 = self.qsrx_band1_int(folder,'BITE')
        
        self.ps_output(0)

        for i in range(4):
            df.iloc[i,4] = p21[i]
            df.iloc[i,5] = p41[i]
            df.iloc[i,6] = max(p21[i],p41[i])
            df.iloc[i,7] = a1[i]
        
        df.to_excel("test_data/%s/%s_QSRx_integrated_SFB1_BITE.xlsx" % (folder,folder),index_label="Band")
        
        self.report_qsr(folder,"template_BITE.docx","report_QSRx_Integrated_SFB1_BITE.docx","report_QSRx_Integrated_SFB1_BITE.pdf",df)

    def track_qsr_sfb2_bite(self):
        df = self.make_df()
        
        dut = self.sfb2_sl+"_"+self.qsrx_sl

        rm = visa.ResourceManager()
        device = rm.open_resource(self.ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()

        folder = self.new_folder(dut,True,True)
        a2,p22,p42 = self.qsrx_band2_int(folder,'BITE')

        self.ps_output(0)

        for j in range(9):
            df.iloc[4+j,4] = p22[j]
            df.iloc[4+j,5] = p42[j]
            df.iloc[4+j,6] = max(p22[j],p42[j])
            df.iloc[4+j,7] = a2[j]

        df.to_excel("test_data/%s/%s_QSRx_integrated_SFB2_BITE.xlsx" % (folder,folder),index_label="Band")
        
        self.report_qsr(folder,"template_BITE.docx","report_QSRx_Integrated_SFB2_BITE.docx","report_QSRx_Integrated_SFB2_BITE.pdf",df)

    def track_qsr_sfb3_bite(self):
        df = self.make_df()
        
        dut = self.sfb3_sl+"_"+self.qsrx_sl

        rm = visa.ResourceManager()
        device = rm.open_resource(self.ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()

        folder = self.new_folder(dut,True,True)
        a3,p23,p43 = self.qsrx_band3_int(folder,'BITE')
        
        self.ps_output(0)

        for k in range(25):
            df.iloc[13+k,4] = p23[k]
            df.iloc[13+k,5] = p43[k]
            df.iloc[13+k,6] = max(p23[k],p43[k])
            df.iloc[13+k,7] = a3[k]
        
        df.to_excel("test_data/%s/%s_QSRx_integrated_SFB3_BITE.xlsx" % (folder,folder),index_label="Band")
        
        self.report_qsr(folder,"template_BITE.docx","report_QSRx_Integrated_SFB3_BITE.docx","report_QSRx_Integrated_SFB3_BITE.pdf",df)

    def make_df(self):
        qsrx = [[0.5,0.8,14.75,15,0.0,0.0,0.0,0.0],
                [0.7,1.2,14.55,14.6,0.0,0.0,0.0,0.0],
                [1.1,1.6,14.15,13.8,0.0,0.0,0.0,0.0],
                [1.5,2,13.75,13,0.0,0.0,0.0,0.0],
                [2,2.5,13.25,12,0.0,0.0,0.0,0.0],
                [2.5,3,12.75,11,0.0,0.0,0.0,0.0],
                [3,3.5,12.25,10,0.0,0.0,0.0,0.0],
                [3.5,4,11.75,9,0.0,0.0,0.0,0.0],
                [4,4.5,11.25,8,0.0,0.0,0.0,0.0],
                [4.5,5,10.75,7,0.0,0.0,0.0,0.0],
                [4.5,5,16.75,11,0.0,0.0,0.0,0.0],
                [5,5.5,10.25,6,0.0,0.0,0.0,0.0],
                [5.5,6,9.75,5,0.0,0.0,0.0,0.0],
                [6,6.5,9.25,4,0.0,0.0,0.0,0.0],
                [6.5,7,9.75,4,0.0,0.0,0.0,0.0],
                [7,7.5,10.25,4,0.0,0.0,0.0,0.0],
                [7.5,8,10.75,4,0.0,0.0,0.0,0.0],
                [8,8.5,11.25,4,0.0,0.0,0.0,0.0],
                [8.5,9,11.75,4,0.0,0.0,0.0,0.0],
                [9,9.5,12.25,4,0.0,0.0,0.0,0.0],
                [9.5,10,12.75,4,0.0,0.0,0.0,0.0],
                [10,10.5,13.25,4,0.0,0.0,0.0,0.0],
                [10.5,11,13.75,4,0.0,0.0,0.0,0.0],
                [11,11.5,14.25,4,0.0,0.0,0.0,0.0],
                [11.5,12,14.75,4,0.0,0.0,0.0,0.0],
                [12,12.5,9.25,4,0.0,0.0,0.0,0.0],
                [12.5,13,9.75,4,0.0,0.0,0.0,0.0],
                [13,13.5,10.25,4,0.0,0.0,0.0,0.0],
                [13.5,14,10.75,4,0.0,0.0,0.0,0.0],
                [13.5,14,16.75,4,0.0,0.0,0.0,0.0],
                [14,14.5,11.25,4,0.0,0.0,0.0,0.0],
                [14.5,15,11.75,4,0.0,0.0,0.0,0.0],
                [15,15.5,12.25,4,0.0,0.0,0.0,0.0],
                [15.5,16,12.75,4,0.0,0.0,0.0,0.0],
                [16,16.5,13.25,4,0.0,0.0,0.0,0.0],
                [16.5,17,13.75,4,0.0,0.0,0.0,0.0],
                [17,17.5,14.25,4,0.0,0.0,0.0,0.0],
                [17.5,18,14.75,4,0.0,0.0,0.0,0.0]]
        
        df = pd.DataFrame(qsrx,columns = ["Start Freq.(GHz)",
                                    "Stop Freq.(GHz)",
                                    "LO Freq.(GHz)",
                                    "LO2 Freq.(GHz)",
                                    "Phase Tracking - 1,3,4 (Deg)",
                                    "Phase Tracking - 1,2,3 (Deg)",
                                    "Final Phase Tracking (Deg)",
                                    "Amp. Tracking (dB)"],index=range(1,39))
        
        return df

    #HMRx

    def phase_ref_hmr(self,band_num,sub_band,window=2):
        ip_zna = self.ip_zna

        bands = [[[0.4,0.75],[1.2,1.6],[1.6,2.2],[0.4,2.2]],
                [[2.2,18.0],[2.5,18.0],[4.0,18.0],[6.0,18.0]]]
        
        current = bands[band_num-1][sub_band]
        
        self.sfb_settings(current,window=window)
        time.sleep(0.5)
        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)
        device.write(":CALC%s:PAR:SDEF 'Phase%s', 'S21'" % (window,window))
            
        device.write("DISPlay:WINDow%s:STATe ON" % (window)) 
        device.write(":DISP:WIND%s:TITL:DATA 'Phase: Band%s'" % (window,band_num))
        device.write(":DISP:WIND%s:TRAC:EFEED 'Phase%s'" % (window,window))
        device.write(":DISPlay:WINDow%s:TRACe:Y:PDIVision 5,'Phase%s'" % (window,window))
        device.write("DISP:WIND%s:TRAC:Y:RLEV 0, 'Phase%s'" % (window,window))
        device.write("DISP:WIND%s:TRAC:Y:RPOS 50, 'Phase%s'" % (window,window))
        device.write(':CALCulate%s:FORMat %s' % (window,'PHAse'))
        time.sleep(0.5)
        device.write(":TRAC:COPY:MATH 'Phase%s_mem','Phase%s'" % (window,window))
        device.write(":DISP:WIND%s:TRAC:EFEED 'Phase%s_mem'" % (window,window))
        time.sleep(0.5)
        device.write(":CALC%s:MATH:SDEF 'Phase%s / Phase%s_mem'" % (window,window,window))
        device.write(':CALC%s:MATH:STAT ON' % (window))
        time.sleep(0.5)
        device.write("TRAC:COPY:MATH 'Phase%s_Ch1','Phase%s'" % (window,window))
        time.sleep(0.5)
        #device.write(":CALC:MATH:STAT OFF")
        device.write(":DISP:WIND%s:TRAC:EFEED 'Phase%s_Ch1'" % (window,window))
        device.write(":DISPlay:WINDow%s:TRACe:Y:PDIVision 5,'Phase%s_Ch1'" % (window,window))
        device.write("DISP:WIND%s:TRAC:Y:RLEV 0, 'Phase%s_Ch1'" % (window,window))
        device.write("DISP:WIND%s:TRAC:Y:RPOS 50, 'Phase%s_Ch1'" % (window,window))
        time.sleep(0.5)

        device.write("DISP:TRAC:SHOW 'Phase%s_mem', OFF" % (window))
        time.sleep(0.5)
        device.close()
        rm.close()

    def gain_ref_hmr(self,band_num,sub_band,window=2):
        ip_zna = self.ip_zna

        bands = [[[0.4,0.75],[1.2,1.6],[1.6,2.2],[0.4,2.2]],
                [[2.2,18.0],[2.5,18.0],[4.0,18.0],[6.0,18.0]]]
        
        current = bands[band_num-1][sub_band]
        self.sfb_settings(current,window=window)
        time.sleep(0.5)
        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)

        if window==1:
            device.write(":CONFigure:TRACe1:REName 'Trc1', 'Gain%s'" % (window))
        else:
            device.write(":CALC%s:PAR:SDEF 'Gain%s', 'S21'" % (window,window))
            device.write(":DISP:WIND%s:TRAC:EFEED 'Gain%s'" % (window,window))
            device.write(":DISPlay:WINDow%s:TRACe:Y:PDIVision 1,'Gain%s'" % (window,window))
        time.sleep(0.5)
        #device.write("CONFigure:CHANnel2:TRACe:REName 'Gain'")
        device.write("DISP:WIND%s:TRAC:Y:RLEV 0, 'Gain%s'" % (window,window))
        device.write("DISP:WIND%s:TRAC:Y:RPOS 50, 'Gain%s'" % (window,window))
        device.write("DISPlay:WINDow%s:STATe ON" % (window))
        device.write(":DISP:WIND%s:TITL:DATA 'Gain: Band%s'" % (window,band_num))
        device.write(':CALCulate%s:FORMat %s' % (window,'MLOGarithmic'))
        time.sleep(0.5)
        device.write(":TRAC:COPY:MATH 'Gain%s_mem','Gain%s'" % (window,window))
        device.write(":DISP:WIND%s:TRAC:EFEED 'Gain%s_mem'" % (window,window))
        time.sleep(0.5)
        device.write(":CALC%s:MATH:SDEF 'Gain%s / Gain%s_mem'" % (window,window,window))
        device.write(':CALC%s:MATH:STAT ON' % (window))
        time.sleep(0.5)
        device.write("TRAC:COPY:MATH 'Gain%s_Ch1','Gain%s'" % (window,window))
        time.sleep(0.5)
        #device.write(":CALC:MATH:STAT OFF")
        device.write(":DISP:WIND%s:TRAC:EFEED 'Gain%s_Ch1'" % (window,window))
        device.write(":DISPlay:WINDow%s:TRACe:Y:PDIVision 1,'Gain%s_Ch1'" % (window,window))
        device.write("DISP:WIND%s:TRAC:Y:RLEV 0, 'Gain%s_Ch1'" % (window,window))
        device.write("DISP:WIND%s:TRAC:Y:RPOS 50, 'Gain%s_Ch1'" % (window,window))
        device.write("DISP:TRAC:SHOW 'Gain%s_mem', OFF" % (window))
        time.sleep(0.5)
        device.close()
        rm.close()

    def track_hmr1(self):

        rm = visa.ResourceManager()
        device = rm.open_resource(self.ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()

        band_ranges = [" (0.4 GHz to 0.75 GHz)",
                    " (1.2 GHz to 1.6 GHz)",
                    " (1.6 GHz to 2.2 GHz)",
                    " (0.4 GHz to 2.2 GHz)"]
        
        for i in range(4):
            if self.stop_event.is_set():
                return
            band = 'Change to Band' + str(i+1) + band_ranges[i]
            f_band = self.mbox(band)
            if f_band == 'OK':
                self.gain_ref_hmr(band_num=1,sub_band=i,window=2*i+1)
                time.sleep(1)
                self.phase_ref_sfb(band_num=1,sub_band=i,window=2*i+2)
            
        for j in range(3):
            if self.stop_event.is_set():
                return
            mess = 'Change to Channel ' + str(j+2) + '\nChange to Band 1' + band_ranges[0]
            ref_port = self.mbox(mess)
            if ref_port == 'OK':
                self.new_gain(1,j+2)
                time.sleep(1)
                self.new_phase(2,j+2)

                f_band = self.mbox('Change to Band 2 %s' % (band_ranges[1]))
                if f_band == 'OK':
                    self.new_gain(3,j+2)
                    time.sleep(1)
                    self.new_phase(4,j+2)
                
                    f_band = self.mbox('Change to Band 3 %s' % (band_ranges[2]))
                    if f_band == 'OK':
                        self.new_gain(5,j+2)
                        time.sleep(1)
                        self.new_phase(6,j+2)
                    
                        f_band = self.mbox('Change to Band 4 %s' % (band_ranges[3]))
                        if f_band == 'OK':
                            self.new_gain(7,j+2)
                            time.sleep(1)
                            self.new_phase(8,j+2)

        time.sleep(0.5)
        self.marker_hmr(1,False)
        self.ps_output(0)
        dut = self.hmrx_sl
        self.save_diagram("HMRx_low","RF",dut)

    def track_hmr1_bite(self):

        rm = visa.ResourceManager()
        device = rm.open_resource(self.ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()

        band_ranges = [" (0.4 GHz to 0.75 GHz)",
                    " (1.2 GHz to 1.6 GHz)",
                    " (1.6 GHz to 2.2 GHz)",
                    " (0.4 GHz to 2.2 GHz)"]
        
        for i in range(4):
            if self.stop_event.is_set():
                return
            band = 'Change to Band' + str(i+1) + band_ranges[i]
            f_band = self.mbox(band)
            if f_band == 'OK':
                self.gain_ref_hmr(band_num=1,sub_band=i,window=2*i+1)
                time.sleep(1)
                self.phase_ref_sfb(band_num=1,sub_band=i,window=2*i+2)
            
        for j in range(3):
            if self.stop_event.is_set():
                return
            mess = 'Change to Channel ' + str(j+2) + '\nChange to Band 1' + band_ranges[0]
            ref_port = self.mbox(mess)
            if ref_port == 'OK':
                self.new_gain(1,j+2)
                time.sleep(1)
                self.new_phase(2,j+2)

                f_band = self.mbox('Change to Band 2 %s' % (band_ranges[1]))
                if f_band == 'OK':
                    self.new_gain(3,j+2)
                    time.sleep(1)
                    self.new_phase(4,j+2)
                
                    f_band = self.mbox('Change to Band 3 %s' % (band_ranges[2]))
                    if f_band == 'OK':
                        self.new_gain(5,j+2)
                        time.sleep(1)
                        self.new_phase(6,j+2)
                    
                        f_band = self.mbox('Change to Band 4 %s' % (band_ranges[3]))
                        if f_band == 'OK':
                            self.new_gain(7,j+2)
                            time.sleep(1)
                            self.new_phase(8,j+2)

        time.sleep(0.5)
        self.marker_hmr(1,True)
        self.ps_output(0)
        dut = self.hmrx_sl
        self.save_diagram("HMRx_low","BITE",dut)

    def track_hmr2(self):

        rm = visa.ResourceManager()
        device = rm.open_resource(self.ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()

        band_ranges = [" (2.2 GHz to 18.0 GHz)",
                    " (2.5 GHz to 18.0 GHz)",
                    " (4.0 GHz to 18.0 GHz)",
                    " (6.0 GHz to 18.0 GHz)"]
        
        for i in range(4):
            if self.stop_event.is_set():
                return
            band = 'Change to Band' + str(i+1) + band_ranges[i]
            f_band = self.mbox(band)
            if f_band == 'OK':
                self.gain_ref_hmr(band_num=1,sub_band=i,window=2*i+1)
                time.sleep(1)
                self.phase_ref_sfb(band_num=1,sub_band=i,window=2*i+2)
            
        for j in range(3):
            if self.stop_event.is_set():
                return
            mess = 'Change to Channel ' + str(j+2) + '\nChange to Band 1' + band_ranges[0]
            ref_port = self.mbox(mess)
            if ref_port == 'OK':
                self.new_gain(1,j+2)
                time.sleep(1)
                self.new_phase(2,j+2)

                f_band = self.mbox('Change to Band 2 %s' % (band_ranges[1]))
                if f_band == 'OK':
                    self.new_gain(3,j+2)
                    time.sleep(1)
                    self.new_phase(4,j+2)
                
                    f_band = self.mbox('Change to Band 3 %s' % (band_ranges[2]))
                    if f_band == 'OK':
                        self.new_gain(5,j+2)
                        time.sleep(1)
                        self.new_phase(6,j+2)
                    
                        f_band = self.mbox('Change to Band 4 %s' % (band_ranges[3]))
                        if f_band == 'OK':
                            self.new_gain(7,j+2)
                            time.sleep(1)
                            self.new_phase(8,j+2)

        time.sleep(0.5)
        self.marker_hmr(1,False)
        self.ps_output(0)
        dut = self.hmrx_sl
        self.save_diagram("HMRx_high","RF",dut)

    def track_hmr2_bite(self):

        rm = visa.ResourceManager()
        device = rm.open_resource(self.ip_zna)
        time.sleep(2)
        device.write('*RST')
        device.close()
        rm.close()

        band_ranges = [" (2.2 GHz to 18.0 GHz)",
                    " (2.5 GHz to 18.0 GHz)",
                    " (4.0 GHz to 18.0 GHz)",
                    " (6.0 GHz to 18.0 GHz)"]
        
        for i in range(4):
            if self.stop_event.is_set():
                return
            band = 'Change to Band' + str(i+1) + band_ranges[i]
            f_band = self.mbox(band)
            if f_band == 'OK':
                self.gain_ref_hmr(band_num=1,sub_band=i,window=2*i+1)
                time.sleep(1)
                self.phase_ref_sfb(band_num=1,sub_band=i,window=2*i+2)
            
        for j in range(3):
            if self.stop_event.is_set():
                return
            mess = 'Change to Channel ' + str(j+2) + '\nChange to Band 1' + band_ranges[0]
            ref_port = self.mbox(mess)
            if ref_port == 'OK':
                self.new_gain(1,j+2)
                time.sleep(1)
                self.new_phase(2,j+2)

                f_band = self.mbox('Change to Band 2 %s' % (band_ranges[1]))
                if f_band == 'OK':
                    self.new_gain(3,j+2)
                    time.sleep(1)
                    self.new_phase(4,j+2)
                
                    f_band = self.mbox('Change to Band 3 %s' % (band_ranges[2]))
                    if f_band == 'OK':
                        self.new_gain(5,j+2)
                        time.sleep(1)
                        self.new_phase(6,j+2)
                    
                        f_band = self.mbox('Change to Band 4 %s' % (band_ranges[3]))
                        if f_band == 'OK':
                            self.new_gain(7,j+2)
                            time.sleep(1)
                            self.new_phase(8,j+2)

        time.sleep(0.5)
        self.marker_hmr(1,True)
        self.ps_output(0)
        dut = self.hmrx_sl
        self.save_diagram("HMRx_high","BITE",dut)

    def marker_hmr(self,band_num,bite=False):
        ip_zna = self.ip_zna
        rm = visa.ResourceManager()
        device = rm.open_resource(ip_zna)

        #device.write('*RST')
        amp_track = []
        phase_track = []

        for i in [1,3,5,7]:
            for j in range(4):
                device.write(":CALCulate%s:PARameter:SELect 'Gain%s_Ch%s'" % (i,i,j+1))
                device.write(":CALCulate%s:MARKer%s ON" % (i,2*j+1))
                device.write(":CALCulate%s:MARKer%s:MAX" % (i,2*j+1))

                device.write(":CALCulate%s:MARKer%s ON" % (i,2*j+2))
                device.write(":CALCulate%s:MARKer%s:MIN" % (i,2*j+2))

        for m in [2,4,6,8]:
            for n in range(4):
                device.write(":CALCulate%s:PARameter:SELect 'Phase%s_Ch%s'" % (m,m,n+1))
                device.write(":CALCulate%s:MARKer%s ON" % (m,2*n+1))
                device.write(":CALCulate%s:MARKer%s:MAX" % (m,2*n+1))

                device.write(":CALCulate%s:MARKer%s ON" % (m,2*n+2))
                device.write(":CALCulate%s:MARKer%s:MIN" % (m,2*n+2))


        for i in [1,3,5,7]:
            device.write(":CALCulate%s:PARameter:SELect 'Gain%s_Ch1'" % (i,i))
            max_val1 = float(device.query("CALC%s:MARK1:Y?" % (i)))
            min_val1 = float(device.query("CALC%s:MARK2:Y?" % (i)))

            device.write(":CALCulate%s:PARameter:SELect 'Gain%s_Ch2'" % (i,i))
            max_val2 = float(device.query("CALC%s:MARK3:Y?" % (i)))
            min_val2 = float(device.query("CALC%s:MARK4:Y?" % (i)))

            device.write(":CALCulate%s:PARameter:SELect 'Gain%s_Ch3'" % (i,i))
            max_val3 = float(device.query("CALC%s:MARK5:Y?" % (i)))
            min_val3 = float(device.query("CALC%s:MARK6:Y?" % (i)))

            device.write(":CALCulate%s:PARameter:SELect 'Gain%s_Ch4'" % (i,i))
            max_val4 = float(device.query("CALC%s:MARK7:Y?" % (i)))
            min_val4 = float(device.query("CALC%s:MARK8:Y?" % (i)))

            amp = (max([max_val1,max_val2,max_val3,max_val4]) - min([min_val1,min_val2,min_val3,min_val4]))/2
            amp_track.append(amp)

        for j in [2,4,6,8]:
            device.write(":CALCulate%s:PARameter:SELect 'Phase%s_Ch1'" % (j,j))
            max_val1 = float(device.query("CALC%s:MARK1:Y?" % (j)))
            min_val1 = float(device.query("CALC%s:MARK2:Y?" % (j)))

            device.write(":CALCulate%s:PARameter:SELect 'Phase%s_Ch2'" % (j,j))
            max_val2 = float(device.query("CALC%s:MARK3:Y?" % (j)))
            min_val2 = float(device.query("CALC%s:MARK4:Y?" % (j)))

            device.write(":CALCulate%s:PARameter:SELect 'Phase%s_Ch3'" % (j,j))
            max_val3 = float(device.query("CALC%s:MARK5:Y?" % (j)))
            min_val3 = float(device.query("CALC%s:MARK6:Y?" % (j)))

            device.write(":CALCulate%s:PARameter:SELect 'Phase%s_Ch4'" % (j,j))
            max_val4 = float(device.query("CALC%s:MARK7:Y?" % (j)))
            min_val4 = float(device.query("CALC%s:MARK8:Y?" % (j)))

            phase = (max([max_val1,max_val2,max_val3,max_val4]) - min([min_val1,min_val2,min_val3,min_val4]))/2
            phase_track.append(phase)

        self.export_csv_hmr(amp_track,phase_track,band_num,bite)

        device.close()
        rm.close()

    def export_csv_hmr(self,amp_track,phase_track,band_num=1,bite=False):
        now = datetime.now()
        dt_string = now.strftime("%d_%m_%Y_%H_%M")

        bands = [["0.4 - 0.75","1.2 - 1.6","1.6 - 2.2","0.4 - 2.2"],
                ["2.2 - 18.0","2.5 - 18.0","4.0 - 18.0","6.0 - 18.0"]]

        hmr_dict = {"Frequency (GHz)":bands[band_num-1],
                    "Amplitude Tracking (Db)":[0.0,0.0,0.0,0.0],
                    "Phase Tracking (deg)":[0.0,0.0,0.0,0.0]
                    }

        hmr = pd.DataFrame(hmr_dict,index=[1,2,3,4])

        for a in range(4):
            hmr.iloc[a,1] = amp_track[a]
            hmr.iloc[a,2] = phase_track[a]
        
        dut = self.hmrx_sl
        if band_num == 1:
            hmr_band = 'low'
        else:
            hmr_band = 'high'
        if bite:
            filename = "%s_HMRx_%s_BITE_%s.csv" % (dt_string,hmr_band,dut)
        else:
            filename = "%s_HMRx%s_RF_%s.csv" % (dt_string,hmr_band,dut)
        hmr.to_csv(filename,index_label="Band")

    def start_hmr1(self):
        self.stop_event.clear()
        self.task_thread = Thread(target=self.track_hmr1)
        self.task_thread.start()

    def start_hmr1_bite(self):
        self.stop_event.clear()
        self.task_thread = Thread(target=self.track_hmr1_bite)
        self.task_thread.start()

    def start_hmr2(self):
        self.stop_event.clear()
        self.task_thread = Thread(target=self.track_hmr2)
        self.task_thread.start()

    def start_hmr2_bite(self):
        self.stop_event.clear()
        self.task_thread = Thread(target=self.track_hmr2_bite)
        self.task_thread.start()

